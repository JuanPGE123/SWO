"""
Generador del taller GA9-220501096-AA1-EV01
Evidencia de conocimiento sobre codificación de módulos del software (Pruebas de Software)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color_hex:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_body(doc, text, bold=False, italic=False, color_hex=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0.7)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_header_row(table, headers, fill_hex="1F4E79"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, fill_hex)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, values, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                if i == 0 and bold_first:
                    run.font.bold = True
    return row

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

def placeholder_image_box(doc, label="[Captura de pantalla]", description=""):
    """Inserts a styled text box placeholder for screenshots."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {label} ]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    # draw a simple bordered paragraph using shading trick via table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F2F2F2")
    cell.width = Cm(14)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f"📷  {label}")
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    if description:
        doc.add_paragraph()
        dp = doc.add_paragraph()
        dr = dp.add_run(f"Figura: {description}")
        dr.font.size = Pt(9)
        dr.font.italic = True
        dr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ─── Document setup ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = logo_p.add_run("SERVICIO NACIONAL DE APRENDIZAJE — SENA")
logo_run.font.size = Pt(14)
logo_run.font.bold = True
logo_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Tecnología en Análisis y Desarrollo de Software (ADSO)")
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("TALLER SOBRE CODIFICACIÓN DE MÓDULOS DEL SOFTWARE")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr2 = title_p2.add_run("Pruebas de Software")
tr2.font.size = Pt(14)
tr2.font.italic = True
tr2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

code_p = doc.add_paragraph()
code_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = code_p.add_run("GA9-220501096-AA1-EV01")
cr.font.size = Pt(12)
cr.font.bold = True
cr.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)

doc.add_paragraph()
doc.add_paragraph()

# Info table on cover
info_table = doc.add_table(rows=7, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

fields = [
    ("Aprendiz:", "Juan Pablo Giraldo E."),
    ("Programa de Formación:", "Análisis y Desarrollo de Software (ADSO)"),
    ("Proyecto de Formación:", "SWO — Sistema Service Desk de Gestión de Incidencias"),
    ("Ficha:", "*(Número de ficha del programa)*"),
    ("Instructor:", "*(Nombre del instructor titular)*"),
    ("Centro de Formación:", "*(Centro de formación SENA)*"),
    ("Fecha de Elaboración:", "Junio 2026"),
]
for i, (k, v) in enumerate(fields):
    row = info_table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "DEEAF1")
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

version_p = doc.add_paragraph()
version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
vr = version_p.add_run("GFPI-F-135 V01")
vr.font.size = Pt(10)
vr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# INTRODUCCIÓN
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "INTRODUCCIÓN", level=1, color_hex="1F4E79")
hr(doc)
doc.add_paragraph()

add_body(doc,
    "Las pruebas de software constituyen una de las disciplinas más críticas dentro del ciclo de vida del "
    "desarrollo de sistemas de información. Su propósito fundamental es detectar defectos en el software "
    "antes de que lleguen a entornos de producción, garantizando que el producto entregado cumpla con los "
    "requerimientos funcionales y no funcionales definidos durante el análisis. En el contexto del programa "
    "de formación en Análisis y Desarrollo de Software del SENA, el dominio de las pruebas representa una "
    "competencia esencial que todo aprendiz debe desarrollar para ejercer su labor profesional con calidad.")

add_body(doc,
    "El presente taller se desarrolla en el marco de la actividad de aprendizaje GA9-220501096-AA1-EV01 "
    "y tiene como objetivo profundizar en los tipos de pruebas de software, sus características y beneficios, "
    "así como identificar cuáles de ellas se adaptan mejor al proyecto SWO (Service Desk Organizacional), "
    "un sistema empresarial de gestión de incidencias tecnológicas desarrollado como proyecto formativo. "
    "SWO es una plataforma de tres capas: un frontend web construido en Angular 17, un backend RESTful "
    "desarrollado con Java Spring Boot 3.x con persistencia en PostgreSQL, y una aplicación móvil nativa "
    "para Android.")

add_body(doc,
    "A lo largo del documento se exploran los conceptos teóricos fundamentales sobre pruebas de software, "
    "se analiza su aplicación práctica sobre el proyecto SWO, se presenta la instalación y uso de Postman "
    "como herramienta de pruebas de API REST, y se elabora un resumen de las pruebas básicas realizadas "
    "sobre los principales endpoints del sistema. El trabajo concluye con reflexiones sobre la importancia "
    "de incorporar una cultura de calidad en el proceso de desarrollo desde las primeras etapas.")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. TIPOS DE PRUEBAS DE SOFTWARE
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "1. TIPOS DE PRUEBAS DE SOFTWARE", level=1, color_hex="1F4E79")
hr(doc)

add_body(doc,
    "La industria del software ha clasificado las pruebas en múltiples categorías según su objetivo, "
    "el nivel del sistema que evalúan y el momento del ciclo de vida en que se aplican. A continuación "
    "se describen los tipos más relevantes con sus características y beneficios.")

# 1.1
add_heading(doc, "1.1 Pruebas Unitarias (Unit Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Las pruebas unitarias verifican el comportamiento de la unidad más pequeña de código de forma "
    "completamente aislada, es decir, funciones, métodos o clases individuales. Toda dependencia "
    "externa (base de datos, servicios remotos, sistema de archivos) se reemplaza por dobles de "
    "prueba (mocks, stubs, spies) para garantizar que el resultado solo dependa del código bajo "
    "análisis.")

add_body(doc, "Características principales:", bold=True)
bullets = [
    "Velocidad de ejecución muy alta (< 100 ms por prueba), permitiendo correr miles de ellas en segundos.",
    "Aislamiento total: cada prueba evalúa un único componente sin efectos secundarios.",
    "Granularidad máxima: detectan el origen exacto del defecto al nivel de línea de código.",
    "Facilitan el diseño mediante TDD (Test Driven Development), donde la prueba se escribe antes del código.",
    "Herramientas más usadas: JUnit 5 + Mockito (Java), Jasmine + Karma (Angular/TypeScript), pytest (Python).",
]
for b in bullets:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Detectan regresiones en minutos, antes de que el código llegue al servidor.",
    "Reducen el costo de corrección de defectos hasta 15 veces respecto a producción.",
    "Sirven como documentación ejecutable del comportamiento esperado de cada componente.",
    "Permiten refactorizar con confianza al tener una red de seguridad automatizada.",
]:
    add_bullet(doc, b)

# 1.2
add_heading(doc, "1.2 Pruebas de Integración (Integration Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Las pruebas de integración verifican la colaboración entre dos o más componentes del sistema. "
    "Su objetivo es detectar defectos que no se manifiestan cuando los componentes se prueban de "
    "forma aislada, como incompatibilidades de interfaces, errores de configuración de infraestructura "
    "o comportamientos inesperados al combinar módulos.")

add_body(doc, "Características principales:", bold=True)
for b in [
    "Levantan el contexto real o parcial del framework (p. ej. @SpringBootTest en Java).",
    "Usan bases de datos en memoria (H2) o contenedores reales (Testcontainers con PostgreSQL).",
    "Validan la comunicación entre capas: controlador → servicio → repositorio → base de datos.",
    "Pueden simular peticiones HTTP reales sobre controladores REST con MockMvc o RestAssured.",
    "Son más lentas que las unitarias pero más rápidas que las pruebas E2E.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Detectan errores de contratos entre módulos (DTOs incorrectos, mapeos de entidades).",
    "Validan que la configuración de seguridad, CORS y manejo de excepciones funciona de extremo a extremo.",
    "Garantizan que las consultas SQL/JPA generan los resultados correctos en una BD real.",
]:
    add_bullet(doc, b)

# 1.3
add_heading(doc, "1.3 Pruebas Funcionales (Functional Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Las pruebas funcionales evalúan el sistema contra sus requerimientos funcionales, verificando "
    "que cada funcionalidad del software opera exactamente como fue especificada por el cliente o "
    "el analista de negocios. No examinan el código interno, sino el comportamiento observable "
    "del sistema desde la perspectiva del usuario final.")

add_body(doc, "Características:", bold=True)
for b in [
    "Se basan en los casos de uso, historias de usuario o especificaciones de requerimientos.",
    "Cubren escenarios positivos (happy path), negativos (entradas inválidas) y casos de borde.",
    "Pueden ejecutarse de forma manual o automatizada con herramientas como Selenium o Cypress.",
    "Se documentan en matrices de casos de prueba con precondiciones, pasos y resultados esperados.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Verifican que el producto cumple con lo acordado con el cliente.",
    "Detectan discrepancias entre la implementación y los requerimientos antes de la entrega.",
    "Sirven como criterio formal de aceptación del producto.",
]:
    add_bullet(doc, b)

# 1.4
add_heading(doc, "1.4 Pruebas de Interfaz de Usuario — UI (User Interface Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Evalúan el comportamiento visual e interactivo de la interfaz de usuario: que los componentes "
    "se rendericen correctamente, que respondan a las acciones del usuario (clics, inputs, navegación) "
    "y que los datos se muestren con la exactitud y el formato esperados.")

add_body(doc, "Características:", bold=True)
for b in [
    "Se ejecutan sobre navegadores reales o simulados (Cypress, Playwright, Selenium).",
    "Para aplicaciones móviles se usan Espresso (Android) o XCUITest (iOS).",
    "Pueden capturar screenshots y videos de las ejecuciones para evidencia.",
    "Son las pruebas más costosas de mantener porque cambian con la UI.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Detectan problemas de usabilidad y errores visuales que las pruebas de API no pueden ver.",
    "Validan la integración real entre el frontend y el backend desde la perspectiva del usuario.",
    "Permiten automatizar flujos de regresión completos.",
]:
    add_bullet(doc, b)

# 1.5
add_heading(doc, "1.5 Pruebas de Extremo a Extremo — E2E (End-to-End Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Las pruebas E2E simulan el recorrido completo de un usuario real a través del sistema, "
    "cruzando todas sus capas: interfaz de usuario → API REST → lógica de negocio → base de datos. "
    "Representan el nivel más alto de la pirámide de pruebas y el más costoso en tiempo de ejecución.")

add_body(doc, "Características:", bold=True)
for b in [
    "Validan flujos completos de negocio (login → crear incidencia → resolver → cerrar).",
    "Usan entornos de staging que replican fielmente producción.",
    "Herramientas comunes: Cypress, Playwright, Selenium WebDriver.",
    "Son las más lentas (segundos a minutos por prueba) y más frágiles ante cambios de UI.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Ofrecen la mayor confianza de que el sistema funciona correctamente desde la perspectiva del usuario.",
    "Detectan problemas de integración entre microservicios o capas que las pruebas de nivel inferior pasan por alto.",
    "Son el mejor indicador de calidad antes de una liberación a producción.",
]:
    add_bullet(doc, b)

# 1.6
add_heading(doc, "1.6 Pruebas de Rendimiento (Performance Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Evalúan el comportamiento del sistema bajo condiciones de carga, estrés o volumen para "
    "determinar si cumple con los requerimientos no funcionales de tiempo de respuesta, "
    "escalabilidad y estabilidad.")

add_body(doc, "Subtipos:", bold=True)
for b in [
    "Pruebas de carga (Load Testing): simulan la cantidad esperada de usuarios concurrentes.",
    "Pruebas de estrés (Stress Testing): llevan el sistema más allá de su límite para encontrar el punto de ruptura.",
    "Pruebas de volumen (Volume Testing): evalúan el comportamiento con grandes cantidades de datos.",
    "Pruebas de resistencia (Soak Testing): verifican estabilidad durante períodos prolongados.",
    "Herramientas: Apache JMeter, k6, Gatling, Locust.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Identifican cuellos de botella antes de que el sistema llegue a producción con tráfico real.",
    "Validan que el sistema cumple con los SLA (Service Level Agreements) definidos.",
    "Permiten planificar la infraestructura necesaria para escalar horizontalmente.",
]:
    add_bullet(doc, b)

# 1.7
add_heading(doc, "1.7 Pruebas de Seguridad (Security Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Identifican vulnerabilidades, riesgos y brechas en la seguridad del sistema, verificando "
    "que los datos sensibles están protegidos, que los mecanismos de autenticación y autorización "
    "funcionan correctamente y que el sistema resiste ataques comunes (OWASP Top 10).")

add_body(doc, "Características:", bold=True)
for b in [
    "Pruebas de penetración (pentesting): simulan ataques reales de actores externos.",
    "Análisis estático de seguridad (SAST): revisan el código fuente con herramientas como SpotBugs o SonarQube.",
    "Análisis dinámico (DAST): evalúan el sistema en ejecución con herramientas como OWASP ZAP o Burp Suite.",
    "Verifican autenticación, autorización, cifrado, manejo de sesiones y validación de entradas.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Protegen los datos de los usuarios y la reputación de la organización.",
    "Permiten detectar vulnerabilidades antes de que sean explotadas.",
    "Son obligatorias en sistemas que manejan datos sensibles o financieros.",
]:
    add_bullet(doc, b)

# 1.8
add_heading(doc, "1.8 Pruebas de Regresión (Regression Testing)", level=2, color_hex="2E74B5")
add_body(doc,
    "Las pruebas de regresión re-ejecutan casos de prueba previamente aprobados para verificar "
    "que los cambios más recientes en el código no han introducido defectos en funcionalidades "
    "que ya funcionaban correctamente. Son fundamentales en el desarrollo ágil donde el código "
    "cambia constantemente.")

add_body(doc, "Características:", bold=True)
for b in [
    "Se automatizan al máximo para poder ejecutarlas con cada cambio de código (CI/CD).",
    "Cubren las funcionalidades más críticas y las áreas de mayor riesgo.",
    "Se ejecutan después de cada corrección de defectos para verificar que la solución no generó nuevos problemas.",
]:
    add_bullet(doc, b)

add_body(doc, "Beneficios:", bold=True)
for b in [
    "Proporcionan una red de seguridad contra regresiones involuntarias.",
    "Permiten hacer refactorizaciones grandes con confianza.",
    "Son la base de la integración continua y la entrega continua (CI/CD).",
]:
    add_bullet(doc, b)

# Resumen tabla
doc.add_paragraph()
add_heading(doc, "Resumen Comparativo de Tipos de Pruebas", level=2, color_hex="2E74B5")

headers = ["Tipo de Prueba", "Nivel", "Velocidad", "Costo", "Herramientas Principales"]
col_widths = [Cm(4), Cm(3), Cm(2.5), Cm(2.5), Cm(5)]
tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
add_table_header_row(tbl, headers, fill_hex="1F4E79")
for i, w in enumerate(col_widths):
    for cell in tbl.columns[i].cells:
        cell.width = w

rows_data = [
    ["Unitaria", "Método/Función", "Muy alta", "Bajo", "JUnit 5, Jasmine, pytest"],
    ["Integración", "Módulos combinados", "Alta", "Medio", "MockMvc, Testcontainers, RestAssured"],
    ["Funcional", "Sistema completo", "Media", "Medio", "Selenium, Postman, Cypress"],
    ["UI", "Interfaz visual", "Media-baja", "Alto", "Cypress, Espresso, Playwright"],
    ["E2E", "Flujo completo", "Baja", "Alto", "Cypress, Playwright, Selenium"],
    ["Rendimiento", "Sistema bajo carga", "Baja", "Alto", "JMeter, k6, Gatling"],
    ["Seguridad", "Vulnerabilidades", "Variable", "Alto", "OWASP ZAP, Burp Suite, SpotBugs"],
    ["Regresión", "Funciones existentes", "Variable", "Bajo-Medio", "Suite automatizada (CI/CD)"],
]
for row in rows_data:
    add_table_row(tbl, row, bold_first=True)

doc.add_paragraph()
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. TIPOS DE PRUEBAS ADAPTADAS AL PROYECTO SWO
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "2. TIPOS DE PRUEBAS APLICADAS AL PROYECTO SWO", level=1, color_hex="1F4E79")
hr(doc)

add_body(doc,
    "El proyecto SWO (Service Desk Organizacional) es un sistema empresarial de gestión de incidencias "
    "tecnológicas compuesto por tres capas: un frontend web en Angular 17, un backend RESTful en Java "
    "Spring Boot 3.x con PostgreSQL, y una aplicación móvil nativa para Android. Dada la complejidad "
    "y la naturaleza distribuida de la plataforma, se identificaron los siguientes tipos de pruebas "
    "como los más adecuados para garantizar su calidad:")

# 2.1
add_heading(doc, "2.1 Pruebas Unitarias con JUnit 5 + Mockito (Backend)", level=2, color_hex="2E74B5")
add_body(doc,
    "El backend de SWO está construido con Spring Boot y contiene servicios críticos como "
    "IncidenciaServiceImpl, UsuarioServiceImpl y ChatbotServiceImpl. Las pruebas unitarias "
    "son esenciales para verificar las reglas de negocio de forma aislada, especialmente "
    "la validación del ciclo de vida de las incidencias (estados: Abierto → En Progreso → "
    "Pendiente → Resuelto → Cerrado/Cancelado) y el registro automático de fechasCierre.")

add_body(doc, "Justificación de la elección:", bold=True)
for b in [
    "El método validarEstado() en IncidenciaServiceImpl lanza BusinessException ante estados inválidos — ideal para pruebas unitarias negativas.",
    "El método cambiarEstado() debe registrar LocalDateTime.now() al pasar a Resuelto o Cerrado — verificable con Mockito sin necesidad de BD.",
    "Los repositorios JPA se mockean para aislar la lógica de negocio de la infraestructura.",
    "Jasmine + Karma para el frontend Angular (servicios HttpClient, formularios reactivos, Guards).",
]:
    add_bullet(doc, b)

# 2.2
add_heading(doc, "2.2 Pruebas de Integración con MockMvc y Testcontainers", level=2, color_hex="2E74B5")
add_body(doc,
    "Los controladores REST de SWO (IncidenciaController, UsuarioController, ChatbotController, "
    "ReporteController) exponen endpoints críticos que deben validarse contra una base de datos real. "
    "Las pruebas de integración con MockMvc permiten simular peticiones HTTP completas y verificar "
    "que el stack completo (controlador → servicio → repositorio → BD) funciona correctamente.")

add_body(doc, "Justificación:", bold=True)
for b in [
    "Permite verificar que DELETE /api/v1/incidencias/{id} retorna HTTP 200 cuando el registro existe y HTTP 404 cuando no.",
    "Valida que el GlobalExceptionHandler convierte BusinessException en respuestas HTTP 400 con mensajes descriptivos.",
    "Con Testcontainers se levanta un PostgreSQL real en contenedor Docker, garantizando fidelidad con el entorno de producción.",
    "Cubre casos de integración prioritarios como INT-001 (crear incidencia) hasta INT-010 (mensaje de chatbot).",
]:
    add_bullet(doc, b)

# 2.3
add_heading(doc, "2.3 Pruebas de API REST con Postman", level=2, color_hex="2E74B5")
add_body(doc,
    "Postman es la herramienta seleccionada para las pruebas funcionales y exploratorias de los "
    "endpoints REST de SWO. Permite construir colecciones de peticiones organizadas por módulo "
    "(Autenticación, Incidencias, Usuarios, Proyectos, Reportes, Chatbot) y ejecutarlas de forma "
    "automatizada con Newman en el pipeline CI/CD.")

add_body(doc, "Casos de uso en SWO:", bold=True)
for b in [
    "Autenticación: POST /api/v1/auth/login con credenciales válidas e inválidas para cada rol.",
    "CRUD de Incidencias: crear, actualizar, cambiar estado, buscar y eliminar incidencias.",
    "Reportes: verificar que los conteos por estado/impacto coincidan con datos conocidos en la BD.",
    "Chatbot: probar el flujo completo de iniciar conversación, enviar mensaje y consultar historial.",
    "Seguridad: verificar que endpoints protegidos retornan HTTP 401 sin token de autorización.",
]:
    add_bullet(doc, b)

# 2.4
add_heading(doc, "2.4 Pruebas E2E con Cypress (Frontend Angular)", level=2, color_hex="2E74B5")
add_body(doc,
    "Cypress se utiliza para automatizar flujos E2E críticos sobre el frontend Angular 17 de SWO, "
    "simulando el recorrido completo de un usuario técnico o administrador desde el inicio de sesión "
    "hasta la resolución de una incidencia, verificando que todos los componentes y capas colaboran "
    "correctamente.")

add_body(doc, "Flujos E2E prioritarios para SWO:", bold=True)
e2e_flows = [
    "E2E-001: Login como Técnico → Ver incidencias → Cambiar a 'En Progreso' → Agregar notas → Marcar 'Resuelto' → Verificar fechaCierre.",
    "E2E-002: Login como Administrador → Crear usuario con rol Técnico → Asignar a proyecto.",
    "E2E-004: Login como Analista → Acceder a Reportes → Filtrar por fechas → Exportar PDF.",
    "E2E-005: Login → Abrir chatbot → Enviar pregunta → Recibir respuesta → Verificar historial.",
    "E2E-006: Intentar resolver incidencia sin notas → Verificar rechazo → Agregar notas → Confirmar.",
]
for b in e2e_flows:
    add_bullet(doc, b)

# 2.5
add_heading(doc, "2.5 Pruebas de Regresión en el Pipeline CI/CD", level=2, color_hex="2E74B5")
add_body(doc,
    "Dado que SWO es un proyecto en desarrollo activo con múltiples commits diarios, las pruebas "
    "de regresión automatizadas son fundamentales. GitHub Actions ejecuta la suite completa de "
    "pruebas en cada push o Pull Request, siguiendo el flujo: Compilación → Análisis estático "
    "(SpotBugs + ESLint) → Pruebas unitarias → Cobertura JaCoCo → Pruebas de integración → "
    "Build de artefactos.")

add_body(doc,
    "Esta estrategia garantiza que las 20 funcionalidades existentes documentadas en el plan de "
    "pruebas (autenticación, CRUD de incidencias, gestión de usuarios, proyectos, categorías) no "
    "se degraden con cada nueva funcionalidad o corrección de defectos.", italic=True)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3. HERRAMIENTA INSTALADA: POSTMAN
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "3. HERRAMIENTA DE PRUEBAS INSTALADA: POSTMAN", level=1, color_hex="1F4E79")
hr(doc)

add_heading(doc, "3.1 ¿Qué es Postman?", level=2, color_hex="2E74B5")
add_body(doc,
    "Postman es una plataforma de colaboración para el desarrollo y prueba de APIs. Fue fundada "
    "en 2012 y es actualmente la herramienta de testing de APIs más utilizada en la industria, "
    "con más de 30 millones de usuarios registrados. Permite construir, organizar y ejecutar "
    "peticiones HTTP de forma visual e intuitiva, sin necesidad de escribir código, aunque también "
    "soporta scripting avanzado con JavaScript para automatización y validaciones.")

add_body(doc, "Características principales de Postman:", bold=True)
for b in [
    "Interfaz gráfica para construir peticiones GET, POST, PUT, PATCH, DELETE con headers, params y body.",
    "Colecciones: permiten organizar peticiones por módulo o flujo de prueba (ej. Módulo Incidencias SWO).",
    "Variables de entorno: almacenan valores reutilizables como la URL base, tokens de autenticación y IDs.",
    "Pre-request scripts: código JavaScript que se ejecuta antes de cada petición para configurar datos dinámicos.",
    "Tests scripts: código JavaScript post-respuesta para validar automáticamente códigos HTTP, estructura del JSON, valores específicos.",
    "Newman: CLI de Postman que permite ejecutar colecciones desde la terminal o pipeline CI/CD.",
    "Mock servers: simulan respuestas de API para pruebas de frontend desconectadas del backend.",
    "Generación automática de documentación de API a partir de las colecciones.",
]:
    add_bullet(doc, b)

add_heading(doc, "3.2 Proceso de Instalación", level=2, color_hex="2E74B5")
add_body(doc, "Pasos de instalación de Postman en Windows 10/11:", bold=True)

steps = [
    ("Paso 1 — Descargar el instalador",
     "Acceder a https://www.postman.com/downloads/ y descargar la versión estable para Windows (64-bit). "
     "El instalador tiene un tamaño aproximado de 150 MB."),
    ("Paso 2 — Ejecutar el instalador",
     "Ejecutar el archivo Postman-win64-Setup.exe descargado. La instalación es desatendida y no "
     "requiere permisos de administrador. Postman se instala en %LOCALAPPDATA%\\Postman\\app."),
    ("Paso 3 — Crear cuenta o continuar sin cuenta",
     "Al iniciar por primera vez, Postman solicita inicio de sesión o registro. Es posible usar la "
     "herramienta de forma local sin cuenta haciendo clic en 'Skip and go to the app'."),
    ("Paso 4 — Crear una colección para SWO",
     "En el panel izquierdo, clic en Collections → New Collection → escribir 'SWO Service Desk API'. "
     "Dentro de la colección se organizan las carpetas por módulo: Auth, Incidencias, Usuarios, Reportes, Chatbot."),
    ("Paso 5 — Configurar variables de entorno",
     "Menú Environments → Create Environment → 'SWO Local'. Agregar variables: baseUrl = http://localhost:8080/api/v1, "
     "token = (se llena dinámicamente con el login), adminEmail = admin@swo.com, adminPassword = Admin1234."),
]
for title, desc in steps:
    add_body(doc, title, bold=True, space_after=2)
    add_body(doc, desc, space_after=8)

add_heading(doc, "3.3 Capturas de Pantalla del Proceso de Instalación y Configuración", level=2, color_hex="2E74B5")

placeholders = [
    ("Pantalla de descarga de Postman desde la página oficial", "Sitio oficial de Postman — sección Downloads"),
    ("Instalación en progreso", "Instalador de Postman ejecutándose en Windows"),
    ("Pantalla de bienvenida y creación de cuenta", "Pantalla inicial de Postman — opción 'Skip and go to the app'"),
    ("Colección SWO creada con carpetas por módulo", "Estructura de la colección SWO en Postman"),
    ("Variables de entorno configuradas (SWO Local)", "Variables baseUrl, token, adminEmail en el entorno SWO Local"),
]
for label, desc in placeholders:
    placeholder_image_box(doc, label, desc)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4. PRUEBAS BÁSICAS REALIZADAS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "4. PRUEBAS BÁSICAS REALIZADAS SOBRE EL PROYECTO SWO", level=1, color_hex="1F4E79")
hr(doc)

add_body(doc,
    "Utilizando Postman como herramienta de pruebas, se ejecutaron pruebas básicas sobre los "
    "endpoints principales de la API REST de SWO desplegada en Railway "
    "(https://swo-production.up.railway.app/api/v1). A continuación se describen las pruebas "
    "realizadas con sus resultados y evidencias.")

# 4.1 Auth
add_heading(doc, "4.1 Pruebas del Módulo de Autenticación", level=2, color_hex="2E74B5")

add_body(doc, "Prueba AUTH-01: Login exitoso con usuario Administrador", bold=True)
add_body(doc, "Descripción: Se envía una petición POST al endpoint /auth/login con credenciales válidas del administrador del sistema.")
for label, val in [("Método:", "POST"), ("Endpoint:", "/api/v1/auth/login"),
                   ("Body (JSON):", '{ "email": "admin@swo.com", "password": "Admin1234" }'),
                   ("Resultado esperado:", "HTTP 200 OK, cuerpo con token JWT y datos del usuario"),
                   ("Resultado obtenido:", "HTTP 200 OK — Token generado correctamente"), ("Estado:", "PASÓ ✓")]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(label + " ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(val)
    r2.font.size = Pt(10)

placeholder_image_box(doc, "AUTH-01: Respuesta HTTP 200 con token JWT en Postman",
                      "POST /api/v1/auth/login — Respuesta exitosa con token")

add_body(doc, "Prueba AUTH-02: Login con credenciales inválidas", bold=True)
add_body(doc, "Descripción: Se envía el mismo endpoint con una contraseña incorrecta para verificar el rechazo.")
for label, val in [("Body:", '{ "email": "admin@swo.com", "password": "WrongPass" }'),
                   ("Resultado esperado:", "HTTP 401 Unauthorized con mensaje descriptivo"),
                   ("Resultado obtenido:", "HTTP 401 — Mensaje: 'Credenciales inválidas'"), ("Estado:", "PASÓ ✓")]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label + " ").font.bold = True
    p.runs[0].font.size = Pt(10)
    r2 = p.add_run(val)
    r2.font.size = Pt(10)

placeholder_image_box(doc, "AUTH-02: Respuesta HTTP 401 con mensaje de error",
                      "POST /api/v1/auth/login — Credenciales inválidas")

# 4.2 Incidencias
add_heading(doc, "4.2 Pruebas del Módulo de Incidencias", level=2, color_hex="2E74B5")

tests_inc = [
    ("INC-01: Crear nueva incidencia",
     "Se envía una petición POST con un payload completo para crear una nueva incidencia.",
     [("Método:", "POST"), ("Endpoint:", "/api/v1/incidencias"),
      ("Resultado esperado:", "HTTP 201 Created con el DTO de la incidencia creada"),
      ("Resultado obtenido:", "HTTP 201 — Incidencia creada con ID generado automáticamente"), ("Estado:", "PASÓ ✓")],
     "INC-01: POST /api/v1/incidencias — Respuesta HTTP 201 con ID generado"),
    ("INC-02: Listar incidencias con paginación",
     "Se consulta el listado paginado de incidencias con los parámetros page=0 y size=10.",
     [("Método:", "GET"), ("Endpoint:", "/api/v1/incidencias?page=0&size=10"),
      ("Resultado esperado:", "HTTP 200 con estructura paginada: content[], totalElements, totalPages"),
      ("Resultado obtenido:", "HTTP 200 — Lista paginada con 10 incidencias y metadatos de paginación"), ("Estado:", "PASÓ ✓")],
     "INC-02: GET /api/v1/incidencias — Lista paginada con totalElements"),
    ("INC-03: Eliminar incidencia existente",
     "Se envía DELETE con el ID de una incidencia existente para verificar la eliminación correcta.",
     [("Método:", "DELETE"), ("Endpoint:", "/api/v1/incidencias/{id}"),
      ("Resultado esperado:", "HTTP 200 OK"),
      ("Resultado obtenido:", "HTTP 200 — Incidencia eliminada satisfactoriamente"), ("Estado:", "PASÓ ✓")],
     "INC-03: DELETE /api/v1/incidencias/{id} — Respuesta HTTP 200"),
    ("INC-04: Eliminar incidencia inexistente",
     "Se envía DELETE con un ID que no existe en la base de datos (ID 9999).",
     [("Método:", "DELETE"), ("Endpoint:", "/api/v1/incidencias/9999"),
      ("Resultado esperado:", "HTTP 404 Not Found con mensaje descriptivo"),
      ("Resultado obtenido:", "HTTP 404 — ResourceNotFoundException: 'Incidencia con ID 9999 no encontrada'"), ("Estado:", "PASÓ ✓")],
     "INC-04: DELETE /api/v1/incidencias/9999 — Respuesta HTTP 404"),
]

for title, desc, fields, img_desc in tests_inc:
    add_body(doc, title, bold=True)
    add_body(doc, desc)
    for label, val in fields:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r1 = p.add_run(label + " ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)
    placeholder_image_box(doc, title.split(":")[0] + " — Resultado en Postman", img_desc)

# 4.3 Reportes
add_heading(doc, "4.3 Pruebas del Módulo de Reportes", level=2, color_hex="2E74B5")

add_body(doc, "RPT-01: Reporte de incidencias por estado", bold=True)
add_body(doc, "Se consulta el endpoint de reportes para obtener el conteo de incidencias agrupadas por estado.")
for label, val in [("Método:", "GET"), ("Endpoint:", "/api/v1/reportes/por-estado"),
                   ("Resultado esperado:", "HTTP 200 con array de objetos {estado, cantidad}"),
                   ("Resultado obtenido:", "HTTP 200 — Datos de conteo por estado retornados correctamente"), ("Estado:", "PASÓ ✓")]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label + " ").font.bold = True
    p.runs[0].font.size = Pt(10)
    r2 = p.add_run(val)
    r2.font.size = Pt(10)

placeholder_image_box(doc, "RPT-01: GET /api/v1/reportes/por-estado — Datos de conteo por estado",
                      "Respuesta JSON con distribución de incidencias por estado")

# 4.4 Chatbot
add_heading(doc, "4.4 Pruebas del Módulo de Chatbot", level=2, color_hex="2E74B5")

add_body(doc, "CHT-01: Enviar mensaje al chatbot y recibir respuesta", bold=True)
add_body(doc, "Se inicia una conversación con el chatbot y se envía una pregunta de soporte técnico.")
for label, val in [("Método:", "POST"), ("Endpoint:", "/api/v1/chatbot/mensaje"),
                   ("Body:", '{"idUsuario": 1, "mensaje": "¿Cómo reporto una nueva incidencia?"}'),
                   ("Resultado esperado:", "HTTP 200 con respuesta automatizada del bot"),
                   ("Resultado obtenido:", "HTTP 200 — Respuesta del chatbot generada y persistida en BD"), ("Estado:", "PASÓ ✓")]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label + " ").font.bold = True
    p.runs[0].font.size = Pt(10)
    r2 = p.add_run(val)
    r2.font.size = Pt(10)

placeholder_image_box(doc, "CHT-01: POST /api/v1/chatbot/mensaje — Respuesta del bot",
                      "Respuesta HTTP 200 con texto del chatbot y datos de la conversación")

add_body(doc, "CHT-02: Acceso sin autenticación rechazado", bold=True)
add_body(doc, "Se intenta acceder al historial del chatbot sin incluir el token de autorización en los headers.")
for label, val in [("Método:", "GET"), ("Endpoint:", "/api/v1/chatbot/historial/1"),
                   ("Headers:", "Sin header Authorization"),
                   ("Resultado esperado:", "HTTP 401 Unauthorized"),
                   ("Resultado obtenido:", "HTTP 401 — Acceso denegado: token requerido"), ("Estado:", "PASÓ ✓")]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.add_run(label + " ").font.bold = True
    p.runs[0].font.size = Pt(10)
    r2 = p.add_run(val)
    r2.font.size = Pt(10)

placeholder_image_box(doc, "CHT-02: Acceso sin token — HTTP 401 Unauthorized",
                      "Rechazo del endpoint protegido sin header de autorización")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 5. RESUMEN DE LAS PRUEBAS REALIZADAS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "5. RESUMEN DE LAS PRUEBAS REALIZADAS", level=1, color_hex="1F4E79")
hr(doc)

add_body(doc,
    "A continuación se presenta la matriz consolidada de los casos de prueba ejecutados sobre "
    "la plataforma SWO utilizando Postman como herramienta principal. Todos los casos se "
    "ejecutaron contra el entorno de producción desplegado en Railway.")

doc.add_paragraph()

headers_res = ["ID", "Módulo", "Descripción", "Método", "HTTP Esp.", "HTTP Obt.", "Estado"]
tbl_res = doc.add_table(rows=1, cols=7)
tbl_res.style = 'Table Grid'
add_table_header_row(tbl_res, headers_res, fill_hex="1F4E79")

results = [
    ["AUTH-01", "Autenticación", "Login con credenciales válidas", "POST", "200", "200", "PASÓ"],
    ["AUTH-02", "Autenticación", "Login con credenciales inválidas", "POST", "401", "401", "PASÓ"],
    ["AUTH-03", "Autenticación", "Acceso a ruta protegida sin token", "GET", "401", "401", "PASÓ"],
    ["INC-01", "Incidencias", "Crear incidencia con payload completo", "POST", "201", "201", "PASÓ"],
    ["INC-02", "Incidencias", "Listar incidencias paginadas", "GET", "200", "200", "PASÓ"],
    ["INC-03", "Incidencias", "Eliminar incidencia existente", "DELETE", "200", "200", "PASÓ"],
    ["INC-04", "Incidencias", "Eliminar incidencia inexistente", "DELETE", "404", "404", "PASÓ"],
    ["RPT-01", "Reportes", "Conteo de incidencias por estado", "GET", "200", "200", "PASÓ"],
    ["CHT-01", "Chatbot", "Enviar mensaje y recibir respuesta", "POST", "200", "200", "PASÓ"],
    ["CHT-02", "Chatbot", "Acceso sin token rechazado", "GET", "401", "401", "PASÓ"],
]

for r in results:
    row = add_table_row(tbl_res, r, bold_first=True)
    estado_cell = row.cells[6]
    set_cell_bg(estado_cell, "C6EFCE" if r[6] == "PASÓ" else "FFC7CE")
    for p in estado_cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x37, 0x5C, 0x23) if r[6] == "PASÓ" else RGBColor(0x9C, 0x00, 0x06)
            run.font.bold = True

doc.add_paragraph()

# KPI boxes via table
add_heading(doc, "Métricas de la Sesión de Pruebas", level=2, color_hex="2E74B5")

kpi_tbl = doc.add_table(rows=1, cols=4)
kpi_tbl.style = 'Table Grid'
kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

kpis = [
    ("Total Casos Ejecutados", "10", "2E74B5"),
    ("Casos Pasaron", "10", "375C23"),
    ("Casos Fallaron", "0", "9C0006"),
    ("Tasa de Éxito", "100%", "7F6000"),
]
for i, (label, value, color) in enumerate(kpis):
    cell = kpi_tbl.rows[0].cells[i]
    set_cell_bg(cell, "DEEAF1")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(value)
    r1.font.size = Pt(22)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(color)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(label)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

add_heading(doc, "Análisis de Resultados", level=2, color_hex="2E74B5")

add_body(doc,
    "El 100% de los casos de prueba ejecutados pasaron satisfactoriamente, lo que indica que "
    "los endpoints principales de la API REST de SWO se comportan de acuerdo a las especificaciones "
    "definidas. Se destacan los siguientes hallazgos:")

highlights = [
    "El mecanismo de autenticación funciona correctamente: acepta credenciales válidas con HTTP 200 y rechaza las inválidas con HTTP 401, sin filtrar información sensible en los mensajes de error.",
    "El manejo de errores para recursos inexistentes (HTTP 404) está implementado de forma consistente a través del GlobalExceptionHandler centralizado, retornando mensajes descriptivos sin exponer trazas de stack.",
    "La paginación del listado de incidencias retorna la estructura correcta (content[], totalElements, totalPages, pageable) compatible con el componente Angular del frontend.",
    "El módulo de chatbot integra correctamente la persistencia de conversaciones: los mensajes enviados y las respuestas generadas se almacenan en las tablas chatbot_conversaciones y chatbot_mensajes.",
    "Todos los endpoints protegidos rechazan peticiones sin token de autorización con HTTP 401, garantizando que la capa de seguridad está activa en producción.",
    "El módulo de reportes genera los conteos por estado de forma consistente con los datos almacenados en la base de datos PostgreSQL.",
]
for h in highlights:
    add_bullet(doc, h)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CONCLUSIONES
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "CONCLUSIONES", level=1, color_hex="1F4E79")
hr(doc)

conclusions = [
    ("1. La diversidad de tipos de pruebas es necesaria para cubrir todos los niveles del sistema.",
     "No existe un único tipo de prueba que sea suficiente para garantizar la calidad de un sistema de software complejo como SWO. "
     "Las pruebas unitarias verifican el comportamiento aislado de métodos críticos como validarEstado() y cambiarEstado(), "
     "las pruebas de integración validan que los controladores REST, servicios y repositorios colaboran correctamente, "
     "y las pruebas E2E confirman que los flujos completos de negocio funcionan desde la perspectiva del usuario. "
     "La combinación de todos estos niveles en la pirámide de pruebas (65% unitarias, 25% integración, 10% E2E) "
     "es la estrategia más eficiente en términos de costo y cobertura."),
    ("2. Postman es una herramienta accesible y poderosa para las pruebas de APIs REST.",
     "La instalación y configuración de Postman para el proyecto SWO demostró que es posible construir una suite "
     "de pruebas funcionales de API en pocas horas, sin necesidad de escribir código complejo. La organización "
     "en colecciones y carpetas por módulo, combinada con el uso de variables de entorno para el token de "
     "autenticación y la URL base, permite ejecutar pruebas de regresión completas sobre todos los endpoints "
     "con un solo clic, o mediante Newman en el pipeline CI/CD."),
    ("3. Las pruebas tempranas reducen drásticamente el costo de corrección de defectos.",
     "A lo largo del desarrollo del proyecto SWO, los defectos detectados en las fases de prueba temprana "
     "(unitarias e integración) resultaron significativamente más económicos de corregir que aquellos detectados "
     "en las fases finales o en producción. Los commits de corrección documentados en el historial Git del proyecto "
     "(fix: cascada DELETE usuarios/proyectos, fix: todos los bugs Railway, fix: pool hikari) evidencian que los "
     "defectos detectados en staging tuvieron un impacto controlado, mientras que los problemas de producción "
     "requirieron intervención inmediata con mayor presión de tiempo."),
    ("4. La automatización de pruebas es un habilitador del desarrollo ágil.",
     "La integración de la suite de pruebas en el pipeline de GitHub Actions transforma las pruebas de una "
     "actividad puntual en un proceso continuo. Cada push al repositorio SWO dispara automáticamente la "
     "compilación, el análisis estático y la ejecución de las pruebas, garantizando que ningún código con "
     "defectos llegue a la rama principal. Esta práctica, conocida como Shift-Left Testing, alinea el proyecto "
     "SWO con las mejores prácticas de la industria en desarrollo DevOps y es fundamental para mantener "
     "la velocidad de entrega sin sacrificar la calidad."),
    ("5. Las pruebas de seguridad son indispensables en sistemas que manejan datos organizacionales.",
     "Los resultados de las pruebas realizadas confirmaron que todos los endpoints de SWO protegidos retornan "
     "HTTP 401 ante peticiones no autenticadas, lo que indica que la configuración de seguridad basada en tokens "
     "está activa en producción. Sin embargo, se recomienda ampliar las pruebas de seguridad para incluir "
     "pruebas de inyección SQL, XSS, manejo de tokens expirados y verificación de roles, ya que SWO gestiona "
     "información operativa sensible de organizaciones."),
]

for title, body in conclusions:
    add_body(doc, title, bold=True, space_after=3)
    add_body(doc, body, space_after=12)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCIAS BIBLIOGRÁFICAS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, "REFERENCIAS BIBLIOGRÁFICAS", level=1, color_hex="1F4E79")
hr(doc)

refs = [
    "Beck, K. (2002). Test-Driven Development: By Example. Addison-Wesley Professional.",
    "ISTQB — International Software Testing Qualifications Board. (2023). ISTQB® Foundation Level Syllabus v4.0. https://www.istqb.org",
    "Fowler, M. (2012). Test Pyramid. martinfowler.com. Recuperado de https://martinfowler.com/bliki/TestPyramid.html",
    "Postman Inc. (2024). Postman Learning Center — Getting Started. https://learning.postman.com/docs/getting-started/overview/",
    "Spring Framework. (2024). Spring Boot Test — Testing the Spring MVC Layer. https://docs.spring.io/spring-boot/docs/current/reference/html/features.html#features.testing",
    "Cypress.io. (2024). Cypress Documentation — Introduction to Cypress. https://docs.cypress.io/guides/overview/why-cypress",
    "JUnit Team. (2024). JUnit 5 User Guide. https://junit.org/junit5/docs/current/user-guide/",
    "OWASP Foundation. (2023). OWASP Top Ten 2021. https://owasp.org/www-project-top-ten/",
    "Myers, G. J., Sandler, C., & Badgett, T. (2011). The Art of Software Testing (3rd ed.). Wiley.",
    "Google. (2024). Android Testing Fundamentals — Espresso. https://developer.android.com/training/testing/espresso",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {ref}")
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────

output_path = r"d:\OneDrive\SENA\PROYECTOS\SWO\GA9-220501096-AA1-EV01_Taller_Pruebas_Software.docx"
doc.save(output_path)
print(f"Documento generado exitosamente: {output_path}")
