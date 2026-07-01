"""
Generador del documento Word: Plan de Validación de Características Mínimas de Hardware
Autor: Script generado para evidencia SENA - ADSO
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
#  HELPERS DE FORMATO
# ─────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    """Rellena el fondo de una celda con color hexadecimal (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="1F3864", sz="4"):
    """Agrega bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    sides = []
    if top:    sides.append("top")
    if bottom: sides.append("bottom")
    if left:   sides.append("left")
    if right:  sides.append("right")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_table_style(table):
    """Aplica estilo base a toda la tabla."""
    table.style = "Table Grid"
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


def header_row(table, cols, bg="1F3864", text_color="FFFFFF", font_size=9):
    """Rellena la fila 0 como encabezado con fondo azul y texto blanco."""
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(text_color)
        set_cell_background(cell, bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def data_row(table, row_idx, values, alt=False, font_size=9, bold_first=False):
    """Escribe una fila de datos con color alternado."""
    bg = "DCE6F1" if alt else "FFFFFF"
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = Pt(font_size)
        if bold_first and i == 0:
            run.bold = True
        set_cell_background(cell, bg)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return p


def add_body(doc, text, bold=False, italic=False, size=10.5, color="2C2C2C", align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, level=0, size=10, color="2C2C2C"):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_numbered(doc, text, size=10, color="2C2C2C"):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_note_box(doc, text, bg="FFF2CC", border_color="F0A500"):
    """Agrega un párrafo con estilo de nota/advertencia."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("  NOTA TÉCNICA:  ")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("7D4E00")
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = RGBColor.from_string("5C3D00")
    return p


def add_formula_block(doc, text):
    """Agrega un bloque de fórmula con fuente monoespaciada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("1A1A2E")
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 95)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string("AAAAAA")


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────
#  DOCUMENTO PRINCIPAL
# ─────────────────────────────────────────────

doc = Document()

# ── Márgenes del documento ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Estilo base de fuente ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


# ════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

# Logo / institución
p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_inst = p_inst.add_run("SERVICIO NACIONAL DE APRENDIZAJE – SENA")
run_inst.bold = True
run_inst.font.size = Pt(13)
run_inst.font.color.rgb = RGBColor.from_string("1F3864")

p_prog = doc.add_paragraph()
p_prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_prog = p_prog.add_run("Tecnología en Análisis y Desarrollo de Software (ADSO)")
run_prog.font.size = Pt(11)
run_prog.font.color.rgb = RGBColor.from_string("2E74B5")

doc.add_paragraph()
doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# Título principal
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run("PLAN DE VALIDACIÓN DE CARACTERÍSTICAS\nMÍNIMAS DE HARDWARE PARA EL DESPLIEGUE\nDE UNA APLICACIÓN WEB")
run_t.bold = True
run_t.font.size = Pt(20)
run_t.font.color.rgb = RGBColor.from_string("1F3864")

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()
doc.add_paragraph()

# Tabla de portada
cover_table = doc.add_table(rows=6, cols=2)
cover_table.style = "Table Grid"
cover_data = [
    ("Nombre del Aprendiz:",   "[Nombre Completo del Aprendiz]"),
    ("Número de Ficha:",       "[XXXXXXX]"),
    ("Nombre del Instructor:", "[Nombre del Instructor]"),
    ("Programa de Formación:", "Análisis y Desarrollo de Software – ADSO"),
    ("Componente Formativo:",  "Infraestructura y Despliegue de Software"),
    ("Fecha de Entrega:",      "Junio de 2026"),
]
for i, (label, value) in enumerate(cover_data):
    row = cover_table.rows[i]
    cell_l = row.cells[0]
    cell_r = row.cells[1]
    cell_l.text = ""
    cell_r.text = ""
    rl = cell_l.paragraphs[0].add_run(label)
    rl.bold = True
    rl.font.size = Pt(10.5)
    rl.font.color.rgb = RGBColor.from_string("FFFFFF")
    rv = cell_r.paragraphs[0].add_run(value)
    rv.font.size = Pt(10.5)
    rv.font.color.rgb = RGBColor.from_string("1F3864")
    set_cell_background(cell_l, "1F3864")
    set_cell_background(cell_r, "EBF3FB")
    set_cell_borders(cell_l)
    set_cell_borders(cell_r)
    cell_l.width = Cm(5.5)
    cell_r.width = Cm(9.5)

doc.add_paragraph()
doc.add_paragraph()

p_version = doc.add_paragraph()
p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
rv = p_version.add_run("Versión 1.0  |  Documento de Evidencia Formativa  |  Uso Académico")
rv.font.size = Pt(9)
rv.italic = True
rv.font.color.rgb = RGBColor.from_string("888888")

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  TABLA DE CONTENIDO (Índice manual)
# ════════════════════════════════════════════════════════════════

add_heading(doc, "TABLA DE CONTENIDO", level=1)
toc_items = [
    ("1.", "Introducción", "3"),
    ("2.", "Lista Ordenada de Elementos de Hardware a Considerar (Checklist de Validación)", "4"),
    ("  2.1", "Procesamiento (CPU)", "4"),
    ("  2.2", "Memoria Volátil (RAM)", "6"),
    ("  2.3", "Almacenamiento (Disco)", "9"),
    ("  2.4", "Conectividad y Red", "12"),
    ("3.", "Análisis de Sistemas Operativos y Licenciamiento de Software", "14"),
    ("  3.1", "Comparativa: Linux vs. Windows Server", "14"),
    ("  3.2", "Modelos de Licenciamiento Open Source vs. Comercial", "17"),
    ("  3.3", "Análisis de Costo Total de Propiedad (TCO)", "19"),
    ("4.", "Dimensionamiento Estimado y Propuesta de Servidor (Sizing)", "21"),
    ("  4.1", "Modelo de Cálculo y Supuestos", "21"),
    ("  4.2", "Especificaciones Técnicas Mínimas — Servidor Físico", "22"),
    ("  4.3", "Especificaciones Técnicas — Instancia Cloud / VPS", "23"),
    ("  4.4", "Configuración Recomendada del Stack de Software", "25"),
    ("5.", "Conclusiones", "27"),
    ("6.", "Fuentes de Referencia (APA 7.ª edición)", "28"),
]
toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.style = "Table Grid"
for i, (num, title, page) in enumerate(toc_items):
    row = toc_table.rows[i]
    row.cells[0].text = num
    row.cells[1].text = title
    row.cells[2].text = page
    bg = "F0F4FA" if i % 2 == 0 else "FFFFFF"
    for j in range(3):
        set_cell_background(row.cells[j], bg)
        set_cell_borders(row.cells[j])
        run = row.cells[j].paragraphs[0].runs
        for r in run:
            r.font.size = Pt(9.5)
    row.cells[0].width = Cm(1.2)
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════

add_heading(doc, "1. INTRODUCCIÓN", level=1)
add_divider(doc)

add_body(doc,
    "El despliegue de una aplicación web en un entorno de producción representa mucho más que la "
    "simple instalación de archivos en un servidor. Constituye una decisión de arquitectura tecnológica "
    "con implicaciones directas sobre la experiencia del usuario, la seguridad de los datos y la "
    "sostenibilidad económica del proyecto. En este contexto, el proceso conocido como Capacity Planning "
    "o Dimensionamiento de Recursos adquiere un carácter crítico e ineludible. El Capacity Planning "
    "es la disciplina de ingeniería que permite cuantificar, proyectar y validar los recursos de hardware "
    "que una aplicación demandará bajo condiciones reales de operación, considerando tanto la carga "
    "presente como el crecimiento futuro esperado. Omitir esta etapa conduce sistemáticamente a dos "
    "escenarios igualmente problemáticos: la subasignación de recursos (underprovisioning), que se "
    "manifiesta como tiempos de respuesta inaceptables, errores 503 Service Unavailable y caídas del "
    "servicio; o la sobreasignación (overprovisioning), que representa un desperdicio del presupuesto "
    "operativo (OPEX) sin ningún beneficio funcional medible para el negocio. Ambos escenarios son "
    "el resultado de la misma causa raíz: la ausencia de un plan formal y técnicamente fundamentado "
    "de validación de hardware previo al despliegue."
)

add_body(doc,
    "En el contexto específico del presente proyecto, se busca desplegar una aplicación web "
    "desarrollada en PHP, operando sobre la pila tecnológica LAMP o LEMP (Linux, Apache/Nginx, "
    "MySQL/PostgreSQL, PHP), con funcionalidades de identificación, autenticación y registro seguro "
    "de usuarios, para una proyección de 250 usuarios activos no concurrentes. Esta configuración "
    "implica que múltiples componentes de software consumirán recursos de hardware de manera simultánea "
    "e interdependiente: el servidor web gestionará las conexiones HTTP/HTTPS entrantes y las balanceará "
    "hacia el intérprete PHP, el motor PHP-FPM procesará la lógica de negocio generando un proceso "
    "o worker por cada petición activa simultánea, y el motor de base de datos relacional mantendrá "
    "en memoria sus propios búferes de caché (InnoDB Buffer Pool) para agilizar las consultas de "
    "autenticación. El presente Plan de Validación tiene como objetivo establecer una metodología "
    "estructurada, una lista de verificación (checklist) exhaustiva y una propuesta de dimensionamiento "
    "técnicamente justificada que garantice que el hardware del servidor sea adecuado, eficiente y "
    "escalable, asegurando un Service Level Agreement (SLA) de disponibilidad superior al 99.5% "
    "desde el primer día de operación en producción."
)

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  2. CHECKLIST DE VALIDACIÓN DE HARDWARE
# ════════════════════════════════════════════════════════════════

add_heading(doc, "2. LISTA ORDENADA DE ELEMENTOS DE HARDWARE A CONSIDERAR", level=1)
add_body(doc,
    "El presente checklist establece los parámetros técnicos, métricas y umbrales mínimos que deben "
    "ser verificados, medidos y documentados en el servidor objetivo antes de proceder con el despliegue "
    "de la aplicación. Cada ítem debe ser comprobado por el responsable técnico, quien deberá registrar "
    "el valor encontrado y firmar su conformidad. Un servidor que no supere los umbrales mínimos en "
    "cualquiera de los ítems críticos no debe recibir cargas de trabajo de producción.",
    size=10.5
)

# ── 2.1 CPU ──────────────────────────────────────────────────────
add_heading(doc, "2.1 Procesamiento (CPU)", level=2, color="2E74B5")

add_body(doc,
    "El procesador es el componente que ejecuta todas las instrucciones del sistema operativo, del "
    "servidor web, del intérprete PHP y del motor de base de datos. Para una aplicación PHP en "
    "producción, la demanda de CPU está directamente correlacionada con la complejidad de los scripts "
    "ejecutados, el número de peticiones concurrentes y la frecuencia de consultas a la base de datos. "
    "PHP-FPM (FastCGI Process Manager) es el gestor de procesos estándar para PHP en producción, y "
    "opera con un modelo de pool de workers: cada worker es un proceso PHP independiente que atiende "
    "una petición a la vez. La capacidad de la CPU determina cuántos workers pueden ejecutarse "
    "eficientemente en paralelo sin generar contención de recursos.",
    size=10.5
)

add_body(doc, "Parámetros críticos a evaluar:", bold=True, size=10.5)
add_bullet(doc, "Arquitectura: Obligatoriamente 64 bits (x86_64/amd64) para soportar más de 4 GB de RAM direccionable y aprovechar el conjunto de instrucciones moderno (SSE4.2, AVX2) que acelera operaciones criptográficas (bcrypt, AES para TLS).")
add_bullet(doc, "Núcleos físicos (Cores): Cada núcleo puede procesar un hilo de ejecución de manera verdaderamente paralela. Con 2 cores, Nginx puede atender conexiones mientras PHP-FPM procesa una petición activa.")
add_bullet(doc, "Hilos lógicos (Threads via Hyper-Threading/SMT): La tecnología de hyper-threading permite que cada núcleo físico gestione 2 hilos lógicos simultáneos, incrementando el throughput en operaciones con tiempo de espera de I/O.")
add_bullet(doc, "Frecuencia de reloj (GHz): Una frecuencia más alta reduce la latencia de ejecución de instrucciones. Para PHP, que es un lenguaje interpretado, una buena frecuencia base es más valiosa que muchos cores.")
add_bullet(doc, "Caché de CPU (L2/L3): La caché integrada almacena instrucciones y datos frecuentemente accedidos, reduciendo la latencia de acceso a RAM. El intérprete PHP se beneficia directamente de una caché L3 amplia.")
add_bullet(doc, "Consumo en estado base (CPU idle): El porcentaje de CPU consumido por el sistema sin carga de aplicación debe ser inferior al 10%. Un valor mayor indica procesos parásitos, malware o hardware degradado.")

doc.add_paragraph()
add_body(doc, "Estimación de workers PHP-FPM para el escenario de 250 usuarios:", bold=True, size=10.5)
add_body(doc,
    "Con 250 usuarios activos distribuidos en una jornada de 8 horas laborales (09:00–17:00), "
    "y considerando que el horario pico concentra el 40% del tráfico en las 2 horas centrales, "
    "la concurrencia máxima estimada se calcula mediante el modelo de Little's Law:",
    size=10.5
)
add_formula_block(doc, "Concurrencia = (Usuarios/hora en pico) × (Duración media de petición PHP en segundos)")
add_formula_block(doc, "Usuarios en hora pico = 250 × 40% / 2h = 50 usuarios/hora")
add_formula_block(doc, "Duración media de petición PHP = 0.5 segundos (operación de autenticación + consulta BD)")
add_formula_block(doc, "Concurrencia máxima = 50 × (0.5 / 3600) ≈ 0.007 → Máx. simultáneos: ~8-15 workers activos")
add_note_box(doc,
    "Aunque la concurrencia matemática máxima estimada es de 8-15 workers, se recomienda configurar "
    "el pool de PHP-FPM con pm.max_children = 25 para absorber picos inesperados y rafagas de tráfico "
    "sin generar errores de 'no available workers'.",
)

doc.add_paragraph()

# Tabla CPU Checklist
add_body(doc, "Checklist de Verificación — CPU:", bold=True, size=10.5)
cpu_cols = ["#", "Parámetro", "Descripción Técnica", "Comando de Verificación", "Umbral Mínimo", "✓/✗"]
cpu_data = [
    ("1", "Arquitectura", "64 bits (x86_64)", "lscpu | grep Architecture", "x86_64", ""),
    ("2", "Cores físicos", "Núcleos de procesamiento real", "lscpu | grep 'Core(s) per socket'", "≥ 2 Cores", ""),
    ("3", "Hilos lógicos", "Hyper-Threading activo", "nproc --all", "≥ 4 Threads", ""),
    ("4", "Frecuencia base", "Clock speed mínima", "lscpu | grep 'CPU MHz'", "≥ 2,000 MHz", ""),
    ("5", "Caché L3", "Memoria caché integrada", "lscpu | grep 'L3 cache'", "≥ 3 MB", ""),
    ("6", "Uso CPU base (idle)", "Sin carga de aplicación", "htop / mpstat 1 5", "< 10%", ""),
    ("7", "Temperatura CPU", "Térmica en operación normal", "sensors (lm-sensors)", "< 75°C", ""),
    ("8", "Vulnerabilidades CPU", "Mitigaciones activas", "cat /sys/devices/system/cpu/vulnerabilities/*", "Mitigated / Not affected", ""),
]
cpu_table = doc.add_table(rows=len(cpu_data)+1, cols=len(cpu_cols))
set_table_style(cpu_table)
header_row(cpu_table, cpu_cols)
for i, row_data in enumerate(cpu_data):
    data_row(cpu_table, i+1, row_data, alt=(i%2==0), font_size=8.5, bold_first=False)

doc.add_paragraph()


# ── 2.2 RAM ──────────────────────────────────────────────────────
add_heading(doc, "2.2 Memoria Volátil (RAM)", level=2, color="2E74B5")

add_body(doc,
    "La memoria RAM es, sin lugar a dudas, el recurso más crítico en entornos PHP. Su subasignación "
    "es la causa más frecuente de degradación del rendimiento en aplicaciones web de producción. "
    "Cuando la RAM física se agota, el kernel Linux comienza a utilizar el espacio de intercambio "
    "(swap space) en disco como memoria virtual. Aunque el swap evita que el sistema colapse "
    "completamente, su latencia de acceso es entre 100 y 1,000 veces mayor que la de la RAM "
    "(nanosegundos vs. milisegundos), lo que se traduce en tiempos de respuesta inaceptables para "
    "el usuario final. En un entorno PHP-FPM, cada worker activo es un proceso independiente del "
    "sistema operativo que consume su propia memoria resident (RSS). Esta característica hace que "
    "el cálculo de RAM sea directo y predecible: la memoria total consumida por PHP-FPM es "
    "proporcional al número máximo de workers configurados, independientemente del tráfico real.",
    size=10.5
)

add_body(doc, "Marco teórico del consumo de memoria por componente:", bold=True, size=10.5)

add_body(doc, "Sistema Operativo (Linux Kernel + Daemons del sistema):", bold=True, size=10, color="1F3864")
add_body(doc,
    "Un servidor Linux en modo headless (sin entorno gráfico), ejecutando únicamente los servicios "
    "esenciales (sshd, systemd, rsyslog, cron), consume entre 200 y 400 MB de RAM. Este consumo "
    "incluye el kernel en memoria, los módulos del kernel cargados, el buffer de red del sistema "
    "y los daemons de administración. Este valor es prácticamente fijo independientemente de la "
    "carga de la aplicación.",
    size=10
)

add_body(doc, "Servidor Web Nginx (Event-Driven Model):", bold=True, size=10, color="1F3864")
add_body(doc,
    "Nginx utiliza un modelo asíncrono no bloqueante basado en eventos (event loop), siguiendo "
    "el patrón del reactor. Esto significa que un número muy reducido de procesos worker (generalmente "
    "igual al número de cores del CPU) gestiona miles de conexiones concurrentes mediante multiplexación "
    "de I/O (epoll en Linux). Cada worker de Nginx consume entre 2 y 8 MB de RAM, resultando en un "
    "consumo total del servidor web de tan solo 50-100 MB para nuestro escenario. Esta eficiencia "
    "es una de las razones principales por las que Nginx desplazó a Apache en entornos de alta "
    "performance.",
    size=10
)

add_body(doc, "PHP-FPM Workers (FastCGI Process Manager):", bold=True, size=10, color="1F3864")
add_body(doc,
    "PHP-FPM mantiene un pool de procesos PHP pre-inicializados listos para atender peticiones. "
    "Cada proceso worker es una instancia completa del intérprete PHP con todas sus extensiones "
    "cargadas en memoria. El consumo de memoria por worker varía significativamente según las "
    "librerías PHP instaladas: una instalación base sin framework consume ~15-20 MB, mientras que "
    "una aplicación con un framework completo (Laravel, Symfony) con ORM y múltiples extensiones "
    "puede consumir 50-80 MB por worker. Para una aplicación personalizada de autenticación de "
    "complejidad media, se estima un consumo de 25-35 MB por worker.",
    size=10
)

add_body(doc, "Motor de Base de Datos MySQL / MariaDB:", bold=True, size=10, color="1F3864")
add_body(doc,
    "MySQL utiliza el parámetro innodb_buffer_pool_size como su principal reserva de memoria. "
    "Este búfer almacena páginas de datos e índices del motor InnoDB en RAM, eliminando la "
    "necesidad de leer desde disco en consultas repetidas. Para una base de datos de usuarios "
    "con 250 registros, el volumen de datos es pequeño y los índices caben completamente en "
    "memoria con tan solo 64-256 MB de buffer pool. Adicionalmente, MySQL reserva memoria para "
    "el buffer de conexiones, el query cache y los logs de transacciones.",
    size=10
)

doc.add_paragraph()
add_body(doc, "Modelo matemático consolidado de consumo de RAM:", bold=True, size=10.5)
add_formula_block(doc, "RAM_TOTAL = RAM_SO + RAM_Nginx + RAM_PHP_FPM + RAM_MySQL + RAM_Reserva_Sistema")
add_formula_block(doc, "RAM_PHP_FPM = pm.max_children × Memoria_por_worker")
add_formula_block(doc, "RAM_PHP_FPM = 25 workers × 30 MB/worker = 750 MB")
add_formula_block(doc, "RAM_TOTAL = 350 MB + 80 MB + 750 MB + 400 MB + 400 MB = ~1,980 MB ≈ 2 GB")
add_formula_block(doc, "RAM Recomendada = RAM_TOTAL × Factor_Seguridad(2.0) = 2 GB × 2 = 4 GB MÍNIMO")
add_note_box(doc,
    "Se recomienda provisionar 4 GB como mínimo absoluto y 8 GB como configuración recomendada para "
    "producción. Los 8 GB proporcionan un margen del 300% sobre el consumo base calculado, absorben "
    "aumentos en el pool de workers, permiten un innodb_buffer_pool_size de 1 GB y dejan memoria "
    "libre para caché del sistema operativo (page cache), que acelera la lectura de archivos PHP."
)

doc.add_paragraph()

# Tabla RAM Checklist
add_body(doc, "Checklist de Verificación — RAM:", bold=True, size=10.5)
ram_cols = ["#", "Parámetro", "Descripción / Propósito", "Comando de Verificación", "Umbral Mínimo", "✓/✗"]
ram_data = [
    ("1", "RAM total física", "Memoria instalada en el servidor", "free -h / dmidecode --type 17", "≥ 4 GB", ""),
    ("2", "RAM disponible (libre + cache)", "Memoria real utilizable para nuevos procesos", "free -h → 'available'", "> 1 GB", ""),
    ("3", "Uso de SWAP activo", "Memoria swap en uso (indica presión de RAM)", "free -h / swapon --show", "< 100 MB (idealmente 0)", ""),
    ("4", "Velocidad de RAM", "Frecuencia de operación de los módulos DRAM", "dmidecode --type 17 | grep Speed", "≥ DDR4 2133 MHz", ""),
    ("5", "Configuración Dual Channel", "Módulos en pares para máximo ancho de banda", "dmidecode --type 17", "2 módulos en pares", ""),
    ("6", "Error ECC (si aplica)", "Memoria con corrección de errores (servidores)", "edac-util -s 4", "No errors reported", ""),
    ("7", "Consumo RSS de PHP-FPM", "Memoria residente por proceso PHP worker", "ps aux | grep php-fpm", "< 40 MB/worker", ""),
    ("8", "OOM Killer activo", "Protección ante agotamiento de memoria", "dmesg | grep -i oom", "Sin eventos OOM recientes", ""),
]
ram_table = doc.add_table(rows=len(ram_data)+1, cols=len(ram_cols))
set_table_style(ram_table)
header_row(ram_table, ram_cols)
for i, row_data in enumerate(ram_data):
    data_row(ram_table, i+1, row_data, alt=(i%2==0), font_size=8.5)

doc.add_paragraph()


# ── 2.3 ALMACENAMIENTO ───────────────────────────────────────────
add_heading(doc, "2.3 Almacenamiento (Disco)", level=2, color="2E74B5")

add_body(doc,
    "El subsistema de almacenamiento es frecuentemente subestimado en los planes de capacidad de "
    "servidores web, pero su impacto en el rendimiento es determinante. En una aplicación PHP, "
    "el disco es accedido en múltiples puntos del ciclo de vida de cada petición: el servidor web "
    "lee los archivos de configuración, PHP-FPM carga los archivos .php desde disco (salvo que "
    "OPcache los tenga en memoria), el motor de base de datos lee y escribe páginas de datos en "
    "sus archivos de datos, y el sistema registra eventos en múltiples archivos de log. La latencia "
    "de cada una de estas operaciones de I/O impacta directamente en el tiempo total de respuesta "
    "de la aplicación (TTFB - Time To First Byte).",
    size=10.5
)

add_body(doc, "Análisis comparativo de tecnologías de almacenamiento:", bold=True, size=10.5)

add_body(doc,
    "HDD (Hard Disk Drive - Disco Duro Mecánico): Los discos magnéticos tradicionales operan "
    "mediante el movimiento físico de cabezales lectores sobre platos giratorios. Esta naturaleza "
    "mecánica introduce una latencia de acceso aleatorio de 5 a 10 milisegundos (tiempo de búsqueda "
    "+ latencia rotacional), que representa 50,000 a 100,000 veces más latencia que la RAM. En un "
    "entorno de base de datos donde MySQL realiza miles de operaciones de I/O aleatorio por minuto, "
    "esta latencia es completamente inaceptable para un entorno de producción.",
    size=10
)

add_body(doc,
    "SSD SATA (Solid State Drive - SATA III): Las unidades de estado sólido eliminan las partes "
    "mecánicas y almacenan datos en memoria flash NAND. La latencia de acceso aleatorio se reduce "
    "a 0.1-0.5 ms, y pueden realizar entre 30,000 y 90,000 IOPS (Input/Output Operations Per Second). "
    "Esto representa una mejora de entre 100 y 1,000 veces respecto al HDD. Un SSD SATA es el "
    "mínimo aceptable para un servidor de base de datos en producción.",
    size=10
)

add_body(doc,
    "NVMe (Non-Volatile Memory Express - PCIe): Las unidades NVMe se conectan directamente al bus "
    "PCIe del procesador, eliminando el cuello de botella del controlador SATA. La latencia cae a "
    "menos de 0.1 ms y los IOPS alcanzan entre 200,000 y 700,000 en unidades de clase enterprise. "
    "El protocolo NVMe fue diseñado específicamente para memoria flash, con colas de comandos "
    "paralelas (hasta 65,535 colas con 65,535 comandos cada una) frente a la única cola del "
    "protocolo AHCI usado por SATA. NVMe es la elección óptima para el directorio de datos de "
    "MySQL (/var/lib/mysql).",
    size=10
)

doc.add_paragraph()

# Tabla comparativa tecnologías disco
add_body(doc, "Tabla comparativa de tecnologías de almacenamiento:", bold=True, size=10.5)
disk_tech_cols = ["Característica", "HDD SATA (7200 RPM)", "SSD SATA III", "NVMe PCIe Gen3", "NVMe PCIe Gen4"]
disk_tech_data = [
    ("Latencia acceso aleatorio", "5 – 10 ms", "0.1 – 0.5 ms", "< 0.1 ms", "< 0.05 ms"),
    ("IOPS (lectura aleatoria 4K)", "80 – 150", "30,000 – 90,000", "300,000 – 500,000", "500,000 – 1,000,000+"),
    ("Velocidad lectura secuencial", "80 – 160 MB/s", "400 – 550 MB/s", "2,000 – 3,500 MB/s", "5,000 – 7,000 MB/s"),
    ("Velocidad escritura secuencial", "80 – 120 MB/s", "300 – 500 MB/s", "1,500 – 3,000 MB/s", "4,000 – 6,500 MB/s"),
    ("Resistencia a vibraciones", "Baja (mecánico)", "Alta (sin partes móv.)", "Alta", "Alta"),
    ("Consumo energético", "5 – 10 W", "2 – 4 W", "3 – 8 W", "4 – 10 W"),
    ("Costo relativo (por GB)", "Bajo ($0.02/GB)", "Medio ($0.08/GB)", "Medio ($0.10/GB)", "Alto ($0.15/GB)"),
    ("Recomendación para producción", "NO ACEPTABLE", "Mínimo aceptable", "Recomendado", "Óptimo"),
]
disk_tech_table = doc.add_table(rows=len(disk_tech_data)+1, cols=len(disk_tech_cols))
set_table_style(disk_tech_table)
header_row(disk_tech_table, disk_tech_cols)
for i, row_data in enumerate(disk_tech_data):
    alt = i % 2 == 0
    row = disk_tech_table.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        bg = "DCE6F1" if alt else "FFFFFF"
        if j == 1 and i == 7: bg = "FFE0E0"
        if j == 2 and i == 7: bg = "FFFACD"
        if j in [3, 4] and i == 7: bg = "D5F5D5"
        set_cell_background(row.cells[j], bg.replace("#",""))
        set_cell_borders(row.cells[j])

doc.add_paragraph()

# Cálculo de espacio en disco
add_body(doc, "Cálculo detallado del espacio en disco requerido:", bold=True, size=10.5)
disk_space_cols = ["Componente", "Tamaño Estimado", "Descripción y Criterio de Cálculo"]
disk_space_data = [
    ("Sistema Operativo (Ubuntu Server 22.04)", "5 – 8 GB", "Instalación base minimal. Incluye kernel, libc, systemd, herramientas esenciales. Sin entorno gráfico."),
    ("Stack LEMP completo", "1 – 2 GB", "Nginx, PHP 8.x + extensiones (pdo, mbstring, openssl, curl, json, opcache), MySQL 8.0, Certbot, Fail2ban, htop, sysstat."),
    ("Código fuente de la aplicación", "500 MB – 2 GB", "Archivos .php, assets estáticos (CSS, JS, imágenes), templates HTML, vendor (dependencias Composer). 2 GB es conservador para una app con muchas librerías."),
    ("Base de datos MySQL (año 1)", "500 MB – 5 GB", "250 usuarios × ~10 KB por registro = 2.5 MB (tablas). Más índices, logs binarios (binlog), undo logs, y datos de sesiones. Factor de crecimiento x20."),
    ("Logs del sistema y aplicación", "5 – 15 GB", "/var/log/nginx/ (access.log, error.log), /var/log/php-fpm/, /var/log/mysql/. Con rotación diaria y retención de 30 días."),
    ("OPcache y archivos temporales de PHP", "256 – 512 MB", "Directorio /tmp para uploads, archivos de sesión PHP (/var/lib/php/sessions/), caché de compilación OPcache en memoria (no en disco)."),
    ("Backups locales temporales", "10 – 30 GB", "Espacio para snapshots de BD antes de migraciones, backups diarios previo a envío a almacenamiento externo (S3, rclone). Retención de 3 días local."),
    ("Margen de crecimiento (24 meses)", "15 – 25 GB", "Factor de crecimiento del 200% proyectado a 2 años. Incluye logs históricos, crecimiento de usuarios y potenciales nuevas funcionalidades."),
    ("TOTAL ESTIMADO", "≈ 37 – 87 GB", ""),
]
disk_space_table = doc.add_table(rows=len(disk_space_data)+1, cols=len(disk_space_cols))
set_table_style(disk_space_table)
header_row(disk_space_table, disk_space_cols)
for i, row_data in enumerate(disk_space_data):
    alt = i % 2 == 0
    row = disk_space_table.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        bg = "DCE6F1" if alt else "FFFFFF"
        if i == len(disk_space_data) - 1:
            bg = "1F3864"
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            run.bold = True
        set_cell_background(row.cells[j], bg.replace("#",""))
        set_cell_borders(row.cells[j])

add_note_box(doc, "Espacio en disco recomendado: Mínimo 80 GB SSD SATA. Para entorno productivo con proyección de 2 años: 120 GB NVMe. Se recomienda separar el volumen de /var/lib/mysql en una partición dedicada para aislar el I/O de la base de datos.")

doc.add_paragraph()

# Checklist disco
add_body(doc, "Checklist de Verificación — Almacenamiento:", bold=True, size=10.5)
disk_check_cols = ["#", "Parámetro", "Comando de Verificación", "Umbral Mínimo", "✓/✗"]
disk_check_data = [
    ("1", "Tecnología del disco (SSD/NVMe)", "lsblk -d -o NAME,ROTA,SIZE,TYPE (ROTA=0 → SSD)", "ROTA = 0", ""),
    ("2", "Espacio total disponible", "df -hT / lsblk -o SIZE", "≥ 80 GB", ""),
    ("3", "Uso actual de particiones clave", "df -h /, df -h /var", "< 70% en uso", ""),
    ("4", "Velocidad de lectura secuencial", "hdparm -Tt /dev/sda (o nvme)", "≥ 300 MB/s", ""),
    ("5", "Salud S.M.A.R.T.", "smartctl -a /dev/sda", "Sin errores críticos, reallocated sectors = 0", ""),
    ("6", "Sistema de archivos", "lsblk -f / df -T", "ext4 o xfs", ""),
    ("7", "IOPS de base de datos", "fio --randread test sobre /var/lib/mysql", "≥ 1,000 IOPS 4K aleatorio", ""),
    ("8", "Montaje con opciones de rendimiento", "cat /etc/fstab", "noatime en partición de datos", ""),
]
disk_check_table = doc.add_table(rows=len(disk_check_data)+1, cols=len(disk_check_cols))
set_table_style(disk_check_table)
header_row(disk_check_table, disk_check_cols)
for i, rd in enumerate(disk_check_data):
    data_row(disk_check_table, i+1, rd, alt=(i%2==0), font_size=8.5)

doc.add_paragraph()


# ── 2.4 RED ──────────────────────────────────────────────────────
add_heading(doc, "2.4 Conectividad y Red", level=2, color="2E74B5")

add_body(doc,
    "La conectividad de red es el canal a través del cual los usuarios interactúan con la aplicación. "
    "Para una aplicación web de autenticación y registro de usuarios, el volumen de datos por "
    "transacción es relativamente bajo comparado con aplicaciones de streaming o transferencia de "
    "archivos, pero la latencia y la disponibilidad del enlace son críticas para la experiencia "
    "del usuario. Un tiempo de respuesta superior a 3 segundos incrementa la tasa de abandono "
    "de usuarios hasta en un 40% (Google/Deloitte Digital, 2018). La red debe garantizar no solo "
    "el ancho de banda suficiente, sino también baja latencia, pérdida mínima de paquetes y "
    "configuración de seguridad adecuada mediante firewalls.",
    size=10.5
)

add_body(doc, "Modelo de cálculo de ancho de banda:", bold=True, size=10.5)
add_body(doc,
    "Para calcular el ancho de banda real requerido, se aplica el siguiente modelo basado en "
    "las características del tráfico:",
    size=10.5
)
add_formula_block(doc, "Supuestos del modelo:")
add_formula_block(doc, "  - Usuarios totales activos en jornada: 250")
add_formula_block(doc, "  - Distribución temporal: jornada de 8 horas, pico en horas 10-12 y 14-16")
add_formula_block(doc, "  - Usuarios simultáneos en hora pico: ~20 usuarios")
add_formula_block(doc, "  - Tamaño promedio de respuesta HTTP (HTML + CSS + JS base): 150 KB")
add_formula_block(doc, "  - Frecuencia de petición: 1 petición cada 30 segundos por usuario activo")
add_formula_block(doc, "")
add_formula_block(doc, "Cálculo de ancho de banda pico (upload desde servidor):")
add_formula_block(doc, "  BW_pico = Usuarios_simultáneos × Tamaño_respuesta × Frecuencia")
add_formula_block(doc, "  BW_pico = 20 usuarios × 150,000 bytes × (1/30 s)")
add_formula_block(doc, "  BW_pico = 100,000 bytes/s = 0.8 Mbps")
add_formula_block(doc, "")
add_formula_block(doc, "Ancho de banda recomendado (con margen de seguridad x10 para picos y overhead TLS):")
add_formula_block(doc, "  BW_recomendado = 0.8 Mbps × 10 = 8 Mbps → Contratar mínimo 10 Mbps simétrico")

doc.add_paragraph()

# Checklist red
add_body(doc, "Checklist de Verificación — Red:", bold=True, size=10.5)
net_cols = ["#", "Parámetro", "Descripción / Propósito", "Comando de Verificación", "Umbral Mínimo", "✓/✗"]
net_data = [
    ("1", "Velocidad de interfaz NIC", "Velocidad máxima de la tarjeta de red", "ethtool eth0 | grep Speed", "≥ 100 Mbps (1 Gbps ideal)", ""),
    ("2", "Ancho de banda contratado (upload)", "Velocidad de subida hacia internet", "speedtest-cli --no-download", "≥ 10 Mbps", ""),
    ("3", "Ancho de banda contratado (download)", "Velocidad de bajada (actualizaciones, S3)", "speedtest-cli --no-upload", "≥ 10 Mbps", ""),
    ("4", "Latencia hacia internet", "Tiempo de ida y vuelta hacia gateway", "ping -c 20 8.8.8.8", "< 50 ms promedio", ""),
    ("5", "Pérdida de paquetes", "Estabilidad del enlace de red", "ping -c 100 8.8.8.8 | tail -2", "< 0.1%", ""),
    ("6", "IP estática/fija asignada", "El servidor debe tener IP fija para DNS", "ip addr show eth0", "IP fija configurada", ""),
    ("7", "Firewall UFW activo", "Protección perimetral del servidor", "ufw status verbose", "Active, puertos 22/80/443 abiertos", ""),
    ("8", "Puertos 80 y 443 abiertos", "HTTP y HTTPS accesibles desde internet", "nc -zv SERVER_IP 80 443", "Connection succeeded", ""),
    ("9", "Fail2ban activo", "Protección anti-brute-force SSH/web", "fail2ban-client status", "Jails activas: sshd, nginx", ""),
    ("10", "DNS resuelto correctamente", "El dominio apunta al IP del servidor", "dig +short dominio.com", "IP del servidor", ""),
]
net_table = doc.add_table(rows=len(net_data)+1, cols=len(net_cols))
set_table_style(net_table)
header_row(net_table, net_cols)
for i, rd in enumerate(net_data):
    data_row(net_table, i+1, rd, alt=(i%2==0), font_size=8.5)

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  3. ANÁLISIS DE S.O. Y LICENCIAMIENTO
# ════════════════════════════════════════════════════════════════

add_heading(doc, "3. ANÁLISIS DE SISTEMAS OPERATIVOS Y LICENCIAMIENTO DE SOFTWARE", level=1)
add_divider(doc)

add_body(doc,
    "La elección del sistema operativo y el modelo de licenciamiento del software de la pila "
    "tecnológica son decisiones de arquitectura con implicaciones técnicas, económicas y "
    "estratégicas de largo plazo. En el contexto de un emprendimiento o proyecto académico que "
    "evoluciona hacia producción, estas decisiones determinan el costo total de propiedad (TCO "
    "— Total Cost of Ownership), la flexibilidad técnica, el nivel de soporte disponible y la "
    "independencia tecnológica del proyecto frente a proveedores específicos (vendor lock-in).",
    size=10.5
)

# ── 3.1 Linux vs Windows ─────────────────────────────────────────
add_heading(doc, "3.1 Comparativa: Linux vs. Windows Server para un Entorno PHP/LAMP", level=2, color="2E74B5")

add_body(doc,
    "La elección del sistema operativo del servidor es la primera y más importante decisión de "
    "infraestructura. A continuación se presenta un análisis técnico multidimensional de las "
    "dos opciones principales del mercado:",
    size=10.5
)

os_cols = ["Criterio de Evaluación", "Linux (Ubuntu 22.04 LTS / Rocky Linux 9)", "Windows Server 2022", "Ventaja"]
os_data = [
    ("Costo de licencia base", "$0 USD — Código Abierto", "$634 – $6,155+ USD (Standard a Datacenter)", "Linux ✓✓"),
    ("Costo CAL (por usuario)", "$0 — No aplica en Linux", "$38-52 USD por usuario. 250 usuarios = $9,500 USD adicionales", "Linux ✓✓"),
    ("Soporte nativo para PHP", "Nativo y óptimo. PHP fue diseñado en y para Linux", "Funcional vía IIS + FastCGI. Requiere configuración adicional", "Linux ✓"),
    ("Rendimiento con Nginx", "Nativo. Nginx usa epoll() de Linux para I/O asíncrono", "Requiere WASI o adaptadores. Menor eficiencia en Windows", "Linux ✓✓"),
    ("Consumo RAM del SO", "200-400 MB en modo headless (sin GUI)", "2-4 GB en modo GUI / ~1.5 GB Server Core", "Linux ✓✓"),
    ("Administración remota", "SSH nativo, CLI. Sin overhead gráfico. Ansible, Bash", "RDP (alto consumo de ancho de banda), PowerShell", "Linux ✓"),
    ("Ciclo de vida LTS", "Ubuntu 22.04: Soporte hasta 2032 (Ubuntu Pro). Rocky 9: 2032", "Windows Server 2022: soporte hasta 2031 (con SA)", "Empate"),
    ("Seguridad (vulnerabilidades)", "Menor superficie de ataque histórica. Parches rápidos (apt)", "Mayor historial de CVEs críticos. Reinicio frecuente para parches", "Linux ✓"),
    ("Contenedores / Docker", "Soporte nativo óptimo. Docker fue diseñado para Linux", "Requiere WSL2 o Hyper-V. Overhead adicional", "Linux ✓✓"),
    ("Integración MySQL/MariaDB", "Nativa. Motor de BD diseñado y optimizado para Linux", "Funcional, pero con menor rendimiento de I/O en disco NTFS", "Linux ✓"),
    ("Automatización / IaC", "Ansible, Terraform, Chef, Puppet — soporte completo nativo", "Soporte parcial en algunas herramientas", "Linux ✓"),
    ("Comunidad de soporte PHP", "Stackoverflow, Ubuntu Forums, RHEL Docs — amplísima", "Menor comunidad para PHP en Windows Server", "Linux ✓"),
    ("Compatibilidad con Cloud", "AMI oficial AWS, imagen oficial GCP/Azure, DigitalOcean", "Disponible pero mayor costo de instancia (licencia incluida)", "Linux ✓"),
    ("Uso global en servidores web", "96.3% del mercado mundial (W3Techs, 2025)", "3.7% del mercado mundial", "Linux ✓✓"),
]
os_table = doc.add_table(rows=len(os_data)+1, cols=len(os_cols))
set_table_style(os_table)
header_row(os_table, os_cols)
for i, row_data in enumerate(os_data):
    alt = i % 2 == 0
    row = os_table.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        bg = "DCE6F1" if alt else "FFFFFF"
        if j == 3 and "Linux" in val:
            bg = "D5F5D5"
        if j == 3 and "Empate" in val:
            bg = "FFFACD"
        set_cell_background(row.cells[j], bg)
        set_cell_borders(row.cells[j])

doc.add_paragraph()

add_body(doc, "Veredicto técnico justificado:", bold=True, size=10.5)
add_body(doc,
    "Para un entorno PHP con stack LAMP/LEMP, Linux es el estándar de la industria sin excepción "
    "técnicamente válida en contra. Los argumentos son contundentes desde múltiples dimensiones:",
    size=10.5
)
add_numbered(doc, "Económica: El ahorro en licencias supera los $11,000 USD en el primer año (ver sección 3.3).")
add_numbered(doc, "Técnica: Linux consume 4-8 veces menos RAM para el SO que Windows Server GUI, dejando significativamente más memoria disponible para los componentes de la aplicación.")
add_numbered(doc, "De rendimiento: Las llamadas al sistema de I/O (epoll, sendfile, splice) disponibles en Linux son las que dan a Nginx su característica eficiencia superior. PHP-FPM, MySQL e incluso los propios benchmarks de PHP son consistentemente 10-30% más rápidos en Linux.")
add_numbered(doc, "De adopción: El 96.3% de los servidores web del mundo corren sobre Linux (W3Techs, 2025). Toda la documentación técnica, tutoriales y guías de hardening están escritos primariamente para Linux.")
add_numbered(doc, "De seguridad: El modelo de permisos POSIX de Linux es más granular y robusto. El menor número de servicios de fondo (background services) en una instalación headless reduce la superficie de ataque.")

add_body(doc, "Distribución Linux recomendada para este proyecto:", bold=True, size=10.5, spacing_after=4)

linux_dist_cols = ["Distribución", "Versión LTS", "Soporte de Seguridad", "Ideal para", "Curva de Aprendizaje"]
linux_dist_data = [
    ("Ubuntu Server 22.04 LTS", "22.04.4 (Jammy)", "Standard: 2027 / ESM (Pro): 2032", "Entornos de aprendizaje, startups, VPS, AWS EC2", "Baja — documentación excelente"),
    ("Rocky Linux 9", "9.4", "2032 (10 años desde lanzamiento)", "Producción empresarial, equivalente a RHEL 9 sin costo", "Media — similar a CentOS"),
    ("Debian 12 (Bookworm)", "Bookworm", "2028 (LTS Debian)", "Servidores muy estables, bajo consumo de recursos", "Media"),
    ("AlmaLinux 9", "9.4", "2032", "Alternativa a Rocky Linux, certificado por CISA", "Media"),
]
ldist_table = doc.add_table(rows=len(linux_dist_data)+1, cols=len(linux_dist_cols))
set_table_style(ldist_table)
header_row(ldist_table, linux_dist_cols)
for i, rd in enumerate(linux_dist_data):
    data_row(ldist_table, i+1, rd, alt=(i%2==0), font_size=8.5)

doc.add_paragraph()
add_note_box(doc, "Para este proyecto académico y las guías del componente formativo SENA, se recomienda Ubuntu Server 22.04 LTS por su documentación exhaustiva, amplia comunidad hispanohablante y excelente integración con herramientas de automatización como Ansible y Docker.")


# ── 3.2 Modelos de Licenciamiento ────────────────────────────────
add_heading(doc, "3.2 Modelos de Licenciamiento de Software: Open Source vs. Comercial", level=2, color="2E74B5")

add_body(doc,
    "El modelo de licenciamiento determina los derechos y obligaciones legales del usuario del "
    "software, así como los costos asociados a su uso en producción. Para el stack LAMP/LEMP, "
    "todos los componentes principales se distribuyen bajo licencias de código abierto (Open Source), "
    "lo que tiene implicaciones técnicas, legales y económicas fundamentales.",
    size=10.5
)

add_body(doc, "Conceptos clave de licenciamiento:", bold=True, size=10.5)
add_body(doc,
    "Licencia de Código Abierto (Open Source): Otorga al usuario cuatro libertades fundamentales "
    "según la Free Software Foundation: (1) ejecutar el software para cualquier propósito, "
    "(2) estudiar y modificar el código fuente, (3) redistribuir copias, y (4) distribuir versiones "
    "modificadas. Dependiendo del tipo específico de licencia (permisiva vs. copyleft), pueden "
    "aplicarse restricciones sobre cómo se redistribuyen las modificaciones.",
    size=10
)
add_body(doc,
    "Licencia Propietaria / Comercial: Otorga únicamente el derecho de uso bajo condiciones "
    "específicas definidas por el titular de los derechos. El código fuente generalmente no es "
    "accesible, la modificación está prohibida y el número de usuarios o instalaciones puede "
    "estar limitado por el número de licencias adquiridas o por Client Access Licenses (CAL).",
    size=10
)
add_body(doc,
    "Client Access Licenses (CAL): Modelo de licenciamiento de Microsoft que requiere una licencia "
    "adicional por cada usuario o dispositivo que accede a los servicios de Windows Server. En un "
    "entorno con 250 usuarios, esto representa 250 CALs adicionales sobre el costo base del SO.",
    size=10
)

doc.add_paragraph()

lic_cols = ["Componente del Stack", "Versión Recomendada", "Tipo de Licencia", "Nombre de Licencia", "Costo por Uso", "Obligaciones Principales"]
lic_data = [
    ("Linux Kernel", "6.x (kernel)", "Copyleft fuerte", "GNU GPL v2", "$0", "Las modificaciones al kernel redistribuidas deben publicarse bajo GPL v2."),
    ("Ubuntu Server 22.04 LTS", "22.04.4 LTS", "Meta-licencia (múltiple)", "Canonical / GPL + MIT + Apache", "$0", "Ubuntu Pro (soporte extendido) es de pago: ~$225/servidor/año."),
    ("Nginx", "1.24 Stable", "Permisiva", "BSD 2-Clause", "$0", "Mantener aviso de copyright en redistribuciones. Sin obligación de publicar modificaciones."),
    ("Apache HTTP Server", "2.4.x", "Permisiva", "Apache License 2.0", "$0", "Incluir aviso NOTICE. Sin obligación de publicar modificaciones. Compatible con GPL."),
    ("PHP Interpreter", "8.2 / 8.3", "Permisiva (tipo BSD)", "PHP License 3.01", "$0", "No usar 'PHP' en el nombre de productos derivados sin permiso. Permite uso en apps propietarias."),
    ("MySQL Community Edition", "8.0 / 8.4", "Copyleft fuerte", "GNU GPL v2", "$0", "Si se distribuye software que enlaza con MySQL (linked), debe ser GPL. Uso en servidor propio es libre."),
    ("MariaDB Community", "10.11 LTS", "Copyleft fuerte", "GNU GPL v2 + LGPL", "$0", "Igual que MySQL CE. La LGPL para librerías cliente permite uso en apps propietarias."),
    ("PostgreSQL", "16.x", "Permisiva", "PostgreSQL License (MIT-style)", "$0", "Una de las licencias más permisivas. Sin restricciones de redistribución o copyleft."),
    ("OpenSSL", "3.x", "Permisiva", "Apache License 2.0", "$0", "Mantener avisos de copyright. Compatible con GPL v3."),
    ("Let's Encrypt / Certbot", "Certbot 2.x", "Permisiva", "Apache License 2.0", "$0", "Certificados SSL/TLS gratuitos con renovación automática cada 90 días."),
    ("Fail2ban", "0.11+", "Copyleft", "GNU GPL v2", "$0", "Uso libre en servidores propios."),
    ("OPcache", "Bundled PHP 8.x", "Permisiva", "PHP License 3.01", "$0", "Incluido y habilitado por defecto en PHP 8.x. No requiere instalación separada."),
]
lic_table = doc.add_table(rows=len(lic_data)+1, cols=len(lic_cols))
set_table_style(lic_table)
header_row(lic_table, lic_cols)
for i, rd in enumerate(lic_data):
    data_row(lic_table, i+1, rd, alt=(i%2==0), font_size=8)

doc.add_paragraph()


# ── 3.3 TCO ──────────────────────────────────────────────────────
add_heading(doc, "3.3 Análisis de Costo Total de Propiedad (TCO) — Primer Año", level=2, color="2E74B5")

add_body(doc,
    "El TCO (Total Cost of Ownership) es una métrica financiera que cuantifica todos los costos "
    "directos e indirectos asociados a la adquisición, implementación, mantenimiento y operación "
    "de un sistema tecnológico durante un período definido. Un análisis de TCO correcto debe "
    "incluir no solo los costos de licenciamiento de software, sino también los costos de hardware, "
    "de infraestructura, de soporte y de administración.",
    size=10.5
)

tco_cols = ["Componente de Costo", "Stack LAMP/LEMP (Open Source)", "Stack Equivalente (Windows + SQL Server)", "Diferencia (Ahorro OSS)"]
tco_data = [
    ("Sistema Operativo", "$0 (Ubuntu Server 22.04 LTS)", "$1,069 USD (Windows Server 2022 Standard 16-core)", "Ahorro: $1,069 USD"),
    ("Client Access Licenses (CAL)", "$0 (No aplica en Linux)", "250 CAL × $38 USD = $9,500 USD", "Ahorro: $9,500 USD"),
    ("Servidor Web", "$0 (Nginx Community)", "IIS — Incluido en Windows Server", "Ahorro: $0"),
    ("Motor de Base de Datos", "$0 (MySQL 8.0 Community Edition)", "SQL Server Standard 2022: $1,418 USD/año", "Ahorro: $1,418 USD"),
    ("Lenguaje PHP", "$0 (PHP 8.3 — Open Source)", "$0 (PHP también corre en Windows)", "Ahorro: $0"),
    ("Certificados SSL/TLS", "$0 (Let's Encrypt — gratuito)", "DigiCert/Comodo Basic SSL: $75-$300 USD/año", "Ahorro: $75-300 USD"),
    ("Soporte de SO (Opcional)", "Ubuntu Pro: $225 USD/año (servidor)", "Software Assurance MS: +30-40% sobre licencia (~$428 USD)", "Ahorro: $203 USD"),
    ("Herramientas de monitoreo", "$0 (Nagios CE, Prometheus, Grafana — OSS)", "$0 (alternativas comerciales: $500-2,000 USD/año)", "Ahorro: $0-2,000 USD"),
    ("Capacitación técnica", "Extensa documentación gratuita, cursos SENA", "Cursos Microsoft Official: $1,500-3,000 USD/persona", "Ahorro: $1,500-3,000 USD"),
    ("TOTAL AÑO 1 (solo licencias)", "$0 – $225 USD", "$12,062 – $13,787 USD", "Ahorro neto: $11,837 – $13,562 USD"),
]
tco_table = doc.add_table(rows=len(tco_data)+1, cols=len(tco_cols))
set_table_style(tco_table)
header_row(tco_table, tco_cols)
for i, row_data in enumerate(tco_data):
    alt = i % 2 == 0
    row = tco_table.rows[i+1]
    is_total = (i == len(tco_data)-1)
    for j, val in enumerate(row_data):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if is_total:
            bg = "1F3864"
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            run.bold = True
        else:
            bg = "DCE6F1" if alt else "FFFFFF"
            if j == 3 and "Ahorro:" in val and "$0" not in val:
                bg = "D5F5D5"
        set_cell_background(row.cells[j], bg)
        set_cell_borders(row.cells[j])

doc.add_paragraph()
add_note_box(doc,
    "IMPACTO EN OPEX: La adopción del stack LAMP/LEMP con componentes 100% Open Source representa "
    "un ahorro de entre $11,837 y $13,562 USD en costos de licenciamiento durante el primer año de "
    "operación. Estos recursos pueden ser reinvertidos en infraestructura de mayor rendimiento, "
    "escalabilidad horizontal, desarrollo de nuevas funcionalidades o en la propia formación del "
    "equipo técnico."
)

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  4. DIMENSIONAMIENTO — SIZING
# ════════════════════════════════════════════════════════════════

add_heading(doc, "4. DIMENSIONAMIENTO ESTIMADO Y PROPUESTA DE SERVIDOR (SIZING)", level=1)
add_divider(doc)

add_body(doc,
    "El dimensionamiento del servidor es la síntesis de todos los análisis previos de hardware. "
    "Su objetivo es traducir los requerimientos de carga de trabajo (workload requirements) en "
    "especificaciones técnicas concretas de hardware. Un dimensionamiento correcto no debe confundirse "
    "con una compra de hardware al límite — debe incluir siempre un factor de seguridad (safety margin) "
    "que absorba variabilidad en el tráfico real, crecimiento futuro de usuarios y el impacto de "
    "actualizaciones de software que puedan incrementar el consumo de recursos.",
    size=10.5
)

# ── 4.1 Modelo de cálculo ─────────────────────────────────────────
add_heading(doc, "4.1 Modelo de Cálculo y Supuestos del Dimensionamiento", level=2, color="2E74B5")

add_body(doc, "Supuestos del modelo de carga de trabajo:", bold=True, size=10.5)
assumptions = [
    ("Usuarios totales activos:", "250 usuarios que acceden a la aplicación durante la jornada laboral."),
    ("Tipo de acceso:", "No concurrente. Los usuarios acceden de forma distribuida, no todos simultáneamente."),
    ("Distribución temporal:", "Jornada de 8 horas (08:00-16:00). Pico de tráfico de 2 horas concentra el 40% de las sesiones."),
    ("Concurrencia estimada en pico:", "15-20 usuarios realizando peticiones activamente al mismo tiempo (calculado por Little's Law)."),
    ("Naturaleza de las peticiones:", "Operaciones de autenticación (SELECT + validación de hash) y registro (INSERT). Complejidad media-baja."),
    ("Stack tecnológico:", "Linux + Nginx + PHP-FPM 8.x + MySQL 8.0 + OPcache habilitado."),
    ("Factor de seguridad aplicado:", "2.0x sobre el consumo base calculado (100% de margen de seguridad)."),
    ("Proyección de crecimiento:", "El sizing debe soportar hasta 500 usuarios activos sin migración de hardware (crecimiento 2x en 18 meses)."),
]
for label, desc in assumptions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run_l = p.add_run(f"{label} ")
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_l.font.color.rgb = RGBColor.from_string("1F3864")
    run_d = p.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = RGBColor.from_string("2C2C2C")

doc.add_paragraph()


# ── 4.2 Servidor físico ───────────────────────────────────────────
add_heading(doc, "4.2 Especificaciones Técnicas — Servidor Físico Dedicado (On-Premise)", level=2, color="2E74B5")

add_body(doc,
    "Un servidor físico dedicado es apropiado cuando la organización requiere control total sobre "
    "el hardware, tiene requisitos de privacidad de datos estrictos, o cuando el análisis de TCO "
    "a 5+ años favorece la inversión de capital (CAPEX) frente al gasto operativo continuo (OPEX) "
    "de la nube. A continuación se presentan las especificaciones mínimas y recomendadas:",
    size=10.5
)

phys_cols = ["Componente", "Especificación Mínima Absoluta", "Especificación Recomendada", "Especificación Óptima (Crecimiento a 2 años)", "Justificación Técnica"]
phys_data = [
    ("Procesador (CPU)",
     "1 socket, 2 cores físicos / 4 threads @ 2.0 GHz\nEj: Intel Xeon E3-1220v5",
     "1 socket, 4 cores / 8 threads @ 3.0+ GHz\nEj: Intel Xeon E-2314 / AMD EPYC 7252",
     "1 socket, 8 cores / 16 threads @ 3.5+ GHz\nEj: Intel Xeon E-2388G / AMD EPYC 7302",
     "4 cores permiten escalar el pool PHP-FPM a 50 workers sin contención de CPU, manteniendo Nginx y MySQL operativos en simultáneo."),
    ("Memoria RAM",
     "4 GB DDR4 2133 MHz (1 módulo)\nMínimo viable, sin margen de crecimiento",
     "8 GB DDR4 3200 MHz (2 × 4 GB Dual Channel)\nConfiguración de producción recomendada",
     "16 GB DDR4 3200 MHz (2 × 8 GB Dual Channel)\nSoporta hasta 500+ usuarios y BD creciente",
     "8 GB proporciona 300% de margen sobre el consumo base de 2 GB calculado, absorbiendo el innodb_buffer_pool en memoria y dejando page cache disponible."),
    ("Almacenamiento Principal",
     "60 GB SSD SATA III\nMínimo para OS + app + logs 1 año",
     "120 GB NVMe PCIe Gen3\nSeparado del SO ideal",
     "240 GB NVMe PCIe Gen4\n+ RAID 1 software (mdadm) o hardware",
     "NVMe reduce la latencia de I/O de MySQL de 0.5 ms (SSD SATA) a < 0.1 ms, acelerando las consultas de autenticación hasta 5x."),
    ("Almacenamiento Secundario",
     "—\n(No requerido en mínimo)",
     "1 TB HDD para backups y logs históricos\nMontado en /backup",
     "2 TB HDD en RAID 1 para backups con retención 30 días",
     "Separar datos de producción (SSD/NVMe) de datos de archivo (HDD) es una práctica de diseño de storage que protege el rendimiento de I/O productivo."),
    ("Interfaz de Red (NIC)",
     "1 × 100 Mbps FastEthernet\nMínimo funcional",
     "1 × 1 Gbps Gigabit Ethernet (RJ-45)\nEstándar de industria actual",
     "2 × 1 Gbps bonded (LACP/802.3ad)\no 1 × 10 Gbps SFP+ para futura escala",
     "1 Gbps es el estándar mínimo en servidores 2023+. La redundancia de NIC (bonding) elimina el NIC como SPOF (Single Point of Failure)."),
    ("Fuente de Poder (PSU)",
     "Simple 300W 80+ Bronze\nPunto único de falla",
     "Simple 400W 80+ Gold\nEficiencia energética mejorada",
     "Redundante 2 × 400W (1+1) 80+ Gold\nElimina SPOF de alimentación",
     "Una PSU redundante garantiza que el fallo de una unidad de alimentación no interrumpa el servicio. Crítico para SLA de alta disponibilidad."),
    ("Sistema de Almacenamiento RAID",
     "Sin RAID (disco único)\nRiesgo total de pérdida de datos ante fallo",
     "RAID 1 software (mdadm) sobre 2 × SSD\nEspejo en tiempo real",
     "RAID 10 hardware sobre 4 × SSD/NVMe\nRendimiento + redundancia",
     "RAID 1 duplica todos los datos en tiempo real entre dos discos. Si uno falla, el servicio continúa sin interrupción. El RAID NO reemplaza los backups."),
    ("Sistema Operativo",
     "Ubuntu Server 22.04.4 LTS (minimal)\nSin GUI, solo paquetes esenciales",
     "Ubuntu Server 22.04.4 LTS\n+ Ubuntu Pro para ESM hasta 2032",
     "Rocky Linux 9.x (Enterprise grade)\nIdeal si el equipo tiene experiencia RHEL",
     "Ubuntu 22.04 LTS es la distribución más documentada para LAMP/LEMP y tiene el mayor ecosistema de soporte comunitario en español."),
]
phys_table = doc.add_table(rows=len(phys_data)+1, cols=len(phys_cols))
set_table_style(phys_table)
header_row(phys_table, phys_cols, font_size=8.5)
for i, row_data in enumerate(phys_data):
    alt = i % 2 == 0
    row = phys_table.rows[i+1]
    for j, val in enumerate(row_data):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if j == 0: run.bold = True
        set_cell_background(row.cells[j], "DCE6F1" if alt else "FFFFFF")
        set_cell_borders(row.cells[j])

doc.add_paragraph()


# ── 4.3 Cloud / VPS ───────────────────────────────────────────────
add_heading(doc, "4.3 Especificaciones Técnicas — Instancia Cloud / VPS", level=2, color="2E74B5")

add_body(doc,
    "Para emprendimientos, proyectos académicos y startups en etapa inicial, el despliegue en "
    "infraestructura cloud o VPS (Virtual Private Server) ofrece ventajas significativas frente "
    "al servidor físico: eliminación de costos de hardware inicial (CAPEX cero), escalabilidad "
    "vertical e horizontal inmediata sin tiempo de inactividad, SLAs de disponibilidad garantizados "
    "por el proveedor (99.95% típico), y costo predictible por modelo de pago por uso. A continuación "
    "se presenta un análisis comparativo de los principales proveedores para este escenario:",
    size=10.5
)

cloud_cols = ["Proveedor", "Tipo de Instancia / Tier", "vCPUs", "RAM", "Almacenamiento", "Ancho de Banda Incluido", "Precio Est./Mes USD", "Recomendación"]
cloud_data = [
    ("AWS EC2", "t3.medium (Mínimo)", "2 vCPU", "4 GB", "30 GB EBS gp3 SSD", "Hasta 5 Gbps (burst)", "~$33", "Viable"),
    ("AWS EC2", "t3.large (Recomendado)", "2 vCPU", "8 GB", "80 GB EBS gp3 SSD", "Hasta 5 Gbps (burst)", "~$60", "✓✓ Recomendado"),
    ("DigitalOcean", "Droplet Basic", "2 vCPU", "4 GB", "80 GB SSD NVMe", "4 TB transferencia", "~$24", "Viable"),
    ("DigitalOcean", "Droplet General (Rec.)", "2 vCPU", "8 GB", "100 GB SSD NVMe", "5 TB transferencia", "~$48", "✓✓ Recomendado"),
    ("Linode/Akamai", "Linode 8GB", "4 vCPU", "8 GB", "160 GB SSD", "5 TB transferencia", "~$48", "✓ Muy buena relación"),
    ("Hetzner Cloud (EU)", "CX31 (Mayor valor)", "2 vCPU", "8 GB", "80 GB SSD NVMe local", "20 TB transferencia", "~€15.90", "✓✓ Mejor valor/precio"),
    ("Vultr", "Regular Cloud Compute", "2 vCPU", "8 GB", "160 GB SSD", "5 TB transferencia", "~$40", "✓ Buena opción"),
    ("Google Cloud", "e2-standard-2", "2 vCPU", "8 GB", "50 GB SSD pd-ssd", "Egress variable", "~$67", "Viable (mayor costo)"),
    ("Azure", "B2s Standard", "2 vCPU", "4 GB", "8 GB SSD temporal", "Variable", "~$35", "Viable con disco extra"),
]
cloud_table = doc.add_table(rows=len(cloud_data)+1, cols=len(cloud_cols))
set_table_style(cloud_table)
header_row(cloud_table, cloud_cols, font_size=8)
for i, rd in enumerate(cloud_data):
    alt = i % 2 == 0
    row = cloud_table.rows[i+1]
    for j, val in enumerate(rd):
        row.cells[j].text = ""
        p = row.cells[j].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(8)
        bg = "DCE6F1" if alt else "FFFFFF"
        if j == 7 and "✓✓" in val: bg = "D5F5D5"
        set_cell_background(row.cells[j], bg)
        set_cell_borders(row.cells[j])

doc.add_paragraph()
add_note_box(doc,
    "Para el contexto de este proyecto (emprendimiento / SENA), la recomendación es DigitalOcean "
    "Droplet General Purpose 8GB o Hetzner Cloud CX31, ambos ofreciendo SSD NVMe local, interfaz "
    "de administración simple y documentación en español. DigitalOcean incluye además una imagen "
    "oficial de LAMP Stack lista para usar."
)

doc.add_paragraph()


# ── 4.4 Configuración del Stack ───────────────────────────────────
add_heading(doc, "4.4 Configuración Recomendada del Stack de Software", level=2, color="2E74B5")

add_body(doc,
    "La correcta parametrización de cada componente del stack es tan importante como el hardware. "
    "Un servidor con especificaciones adecuadas pero con configuración por defecto puede tener "
    "un rendimiento subóptimo. Los siguientes parámetros deben ajustarse según el dimensionamiento "
    "calculado:",
    size=10.5
)

stack_cols = ["Componente", "Parámetro de Configuración", "Valor Recomendado (4 GB RAM / 2 vCPU)", "Justificación"]
stack_data = [
    ("Nginx", "worker_processes", "auto (= número de cores CPU)", "Nginx debe tener un proceso por core para máxima eficiencia."),
    ("Nginx", "worker_connections", "1024", "Define el máximo de conexiones simultáneas por worker. 1024 es más que suficiente para este escenario."),
    ("Nginx", "keepalive_timeout", "65", "Mantiene conexiones HTTP abiertas 65 segundos para reutilización, reduciendo overhead TCP."),
    ("Nginx", "gzip on / gzip_types", "Habilitado para text/html, text/css, application/javascript", "Compresión en tránsito reduce el ancho de banda hasta un 70%."),
    ("PHP-FPM", "pm", "dynamic", "Permite ajuste dinámico del pool de workers según la carga real."),
    ("PHP-FPM", "pm.max_children", "25", "Máximo de workers PHP simultáneos. 25 × 30 MB = 750 MB → dentro del presupuesto de RAM."),
    ("PHP-FPM", "pm.start_servers", "5", "Workers iniciales al arrancar. Reduce latencia en el primer tráfico."),
    ("PHP-FPM", "pm.min_spare_servers", "5", "Mínimo de workers ociosos listos para atender peticiones."),
    ("PHP-FPM", "pm.max_spare_servers", "10", "Máximo de workers ociosos. Libera RAM si el tráfico baja."),
    ("PHP-FPM", "pm.max_requests", "500", "Reinicia un worker después de 500 peticiones para prevenir memory leaks."),
    ("OPcache", "opcache.enable", "1", "Habilita la caché de bytecode PHP. Elimina la necesidad de recompilar scripts en cada petición."),
    ("OPcache", "opcache.memory_consumption", "128 MB", "Memoria RAM dedicada a la caché de bytecode. Suficiente para apps PHP medianas."),
    ("OPcache", "opcache.max_accelerated_files", "10000", "Número máximo de archivos PHP en caché. Adecuado para frameworks modernos."),
    ("OPcache", "opcache.validate_timestamps", "0 (producción)", "En producción, deshabilitar validación de timestamps para máximo rendimiento. Requiere reinicio manual al desplegar código."),
    ("MySQL 8.0", "innodb_buffer_pool_size", "1G (con 4 GB RAM) / 2G (con 8 GB RAM)", "Principal reserva de caché de MySQL. Regla general: 60-70% de la RAM dedicada a MySQL."),
    ("MySQL 8.0", "max_connections", "100", "Máximo de conexiones simultáneas. 100 es generoso para 15-20 usuarios concurrentes."),
    ("MySQL 8.0", "innodb_log_file_size", "256M", "Tamaño del redo log. Un valor mayor mejora el rendimiento de escrituras masivas."),
    ("MySQL 8.0", "query_cache_type", "0 (deshabilitado)", "El query cache de MySQL es contraproducente en MySQL 8.0+ con InnoDB. Usar ProxySQL o Redis para caché de queries si se necesita."),
    ("Sistema (sysctl)", "vm.swappiness", "10", "Reduce la propensión del kernel a usar swap. El sistema preferirá liberar page cache antes de swappear."),
    ("Sistema (sysctl)", "net.core.somaxconn", "1024", "Incrementa el backlog de conexiones TCP pendientes para evitar drops en picos."),
]
stack_table = doc.add_table(rows=len(stack_data)+1, cols=len(stack_cols))
set_table_style(stack_table)
header_row(stack_table, stack_cols)
for i, rd in enumerate(stack_data):
    data_row(stack_table, i+1, rd, alt=(i%2==0), font_size=8.5)

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  5. CONCLUSIONES
# ════════════════════════════════════════════════════════════════

add_heading(doc, "5. CONCLUSIONES", level=1)
add_divider(doc)

add_heading(doc, "5.1 El Capacity Planning como Fundamento de la Alta Disponibilidad", level=3, color="1F3864")
add_body(doc,
    "Un plan de validación de hardware técnicamente riguroso es la base sobre la cual se construye "
    "la alta disponibilidad (High Availability) de cualquier aplicación web. A lo largo de este "
    "documento se demostró que la diferencia entre un servidor subasignado (1 vCPU / 2 GB RAM) y "
    "uno correctamente dimensionado (2 vCPUs / 8 GB RAM) no se mide únicamente en términos de "
    "rendimiento de benchmarks artificiales, sino en la capacidad real del sistema para mantenerse "
    "operativo durante picos de tráfico sin entrar en condiciones de saturación de memoria (swap), "
    "agotamiento del pool de workers PHP-FPM o contención de CPU. Cuando un servidor PHP entra en "
    "swap de memoria, su tiempo de respuesta puede incrementarse de milisegundos a decenas de "
    "segundos, lo que es funcionalmente equivalente a una interrupción del servicio desde la "
    "perspectiva del usuario final. El modelo matemático de consumo presentado en la sección 2.2 "
    "demostró que un servidor con tan solo 4 GB de RAM puede operar cómodamente el stack completo "
    "(Linux + Nginx + PHP-FPM + MySQL) con un margen de seguridad del 100%, y que 8 GB de RAM "
    "ofrece un margen del 300% que absorbe crecimientos inesperados de tráfico y permite incrementar "
    "el innodb_buffer_pool_size de MySQL para maximizar el rendimiento de las consultas de "
    "autenticación. En conclusión, invertir tiempo y rigor técnico en el plan de validación antes "
    "del despliegue previene, con certeza matemática, las causas más frecuentes de indisponibilidad "
    "del servicio, garantizando un SLA de disponibilidad superior al 99.5% (menos de 44 horas de "
    "interrupción acumulada en un año de operación).",
    size=10.5
)

doc.add_paragraph()

add_heading(doc, "5.2 La Elección Estratégica del Stack Open Source como Factor de Sostenibilidad Operativa", level=3, color="1F3864")
add_body(doc,
    "La adopción del stack LAMP/LEMP basado íntegramente en software de código abierto no es "
    "únicamente una decisión técnica, sino una decisión estratégica de sostenibilidad económica "
    "y autonomía tecnológica de largo plazo. Como se cuantificó con precisión en el análisis "
    "comparativo de licenciamiento de la sección 3.3, el ahorro en costos de licencias software "
    "durante el primer año de operación puede superar los $11,837 USD frente a un stack propietario "
    "equivalente basado en Windows Server y SQL Server. Para un emprendimiento o proyecto en etapa "
    "inicial, estos recursos representan la diferencia entre la viabilidad y el estancamiento del "
    "proyecto. Más allá del impacto económico inmediato, la adopción de Linux y del stack open "
    "source proporciona independencia tecnológica total (ausencia de vendor lock-in), lo que "
    "garantiza que el proyecto pueda cambiar de proveedor de infraestructura, de herramienta de "
    "base de datos o de servidor web sin enfrentar costos de migración de licencias ni períodos "
    "de transición forzados por proveedores. La vigencia del soporte de seguridad de Ubuntu 22.04 "
    "LTS hasta el año 2032 (con Ubuntu Pro) garantiza actualizaciones de seguridad por un período "
    "de 10 años desde su lanzamiento, brindando estabilidad y previsibilidad a la plataforma de "
    "despliegue. En conclusión, el plan de validación presentado demuestra que para el escenario "
    "de 250 usuarios activos, un servidor con especificaciones modestas pero correctamente "
    "dimensionadas (2 vCPUs / 8 GB RAM / 80 GB SSD), ejecutando Linux con el stack LAMP/LEMP, "
    "proporciona una plataforma técnicamente sólida, económicamente viable y estratégicamente "
    "sostenible para el despliegue de la aplicación web en producción.",
    size=10.5
)

page_break(doc)


# ════════════════════════════════════════════════════════════════
#  6. FUENTES DE REFERENCIA APA 7
# ════════════════════════════════════════════════════════════════

add_heading(doc, "6. FUENTES DE REFERENCIA", level=1)
add_body(doc, "Formato: APA 7.ª edición", italic=True, size=10, color="555555")
add_divider(doc)

references = [
    "PHP Group. (2024). PHP manual: FastCGI Process Manager (FPM) — Configuration. The PHP Group. https://www.php.net/manual/en/install.fpm.configuration.php",
    "PHP Group. (2024). PHP manual: Zend OPcache. The PHP Group. https://www.php.net/manual/en/book.opcache.php",
    "PHP Group. (2024). PHP license version 3.01. The PHP Group. https://www.php.net/license/3_01.txt",
    "MySQL. (2024). MySQL 8.0 reference manual: InnoDB startup options and system variables — innodb_buffer_pool_size. Oracle Corporation. https://dev.mysql.com/doc/refman/8.0/en/innodb-parameters.html#sysvar_innodb_buffer_pool_size",
    "MySQL. (2024). MySQL 8.0 reference manual: The MySQL server — max_connections. Oracle Corporation. https://dev.mysql.com/doc/refman/8.0/en/server-system-variables.html#sysvar_max_connections",
    "Nginx, Inc. (2024). Nginx documentation: Core functionality (ngx_http_core_module). Nginx, Inc. https://nginx.org/en/docs/http/ngx_http_core_module.html",
    "Nginx, Inc. (2024). Nginx documentation: Optimizing performance for serving content. Nginx, Inc. https://docs.nginx.com/nginx/admin-guide/web-server/serving-static-content/",
    "Canonical Ltd. (2024). Ubuntu server guide: Introductory concepts — Package management. Canonical Ltd. https://ubuntu.com/server/docs",
    "Canonical Ltd. (2024). Ubuntu Pro — security and compliance for Ubuntu. Canonical Ltd. https://ubuntu.com/pro",
    "The Linux Foundation. (2024). Linux Foundation annual report 2024: Open source software statistics. The Linux Foundation. https://www.linuxfoundation.org/research",
    "W3Techs. (2025). Usage statistics of operating systems for websites. Q-Success Web-based Services. https://w3techs.com/technologies/overview/operating_system",
    "Free Software Foundation. (2024). GNU General Public License, version 2 (GPL-2.0). Free Software Foundation. https://www.gnu.org/licenses/old-licenses/gpl-2.0.en.html",
    "Free Software Foundation. (2024). What is free software? The four essential freedoms. Free Software Foundation. https://www.gnu.org/philosophy/free-sw.html",
    "Open Source Initiative. (2024). The Apache License, version 2.0. Open Source Initiative. https://opensource.org/licenses/Apache-2.0",
    "Open Source Initiative. (2024). The BSD 2-Clause License. Open Source Initiative. https://opensource.org/licenses/BSD-2-Clause",
    "Amazon Web Services. (2024). Amazon EC2 instance types: T3 instances — General purpose. Amazon Web Services, Inc. https://aws.amazon.com/ec2/instance-types/t3/",
    "DigitalOcean. (2024). Droplet documentation: Choosing the right Droplet plan for your use case. DigitalOcean, LLC. https://docs.digitalocean.com/products/droplets/",
    "Hetzner Online GmbH. (2024). Cloud server products: CX series. Hetzner Online GmbH. https://www.hetzner.com/cloud",
    "Microsoft Corporation. (2024). Windows Server 2022 licensing datasheet. Microsoft Corporation. https://www.microsoft.com/en-us/licensing/product-licensing/windows-server",
    "Microsoft Corporation. (2024). Client Access License (CAL) overview — Windows Server. Microsoft Corporation. https://www.microsoft.com/en-us/licensing/product-licensing/client-access-license",
    "MariaDB Foundation. (2024). MariaDB server documentation: System variables. MariaDB Foundation. https://mariadb.com/kb/en/server-system-variables/",
    "The PostgreSQL Global Development Group. (2024). PostgreSQL 16 documentation: Server configuration. The PostgreSQL Global Development Group. https://www.postgresql.org/docs/16/runtime-config.html",
    "Internet Security Research Group (ISRG). (2024). Let's Encrypt documentation: Getting started. ISRG. https://letsencrypt.org/getting-started/",
    "Certbot Contributors. (2024). Certbot instructions for Nginx on Ubuntu 22.04. Electronic Frontier Foundation. https://certbot.eff.org/instructions?certbot_apache=n&os=ubuntufocal",
    "Gregg, B. (2020). Systems performance: Enterprise and the cloud (2nd ed.). Addison-Wesley Professional.",
    "Limoncelli, T. A., Hogan, C. J., & Chalup, S. R. (2022). The practice of system and network administration (3rd ed.). Addison-Wesley Professional.",
    "Google & Deloitte Digital. (2018). The need for mobile speed: How mobile latency impacts publisher revenue. Think with Google. https://www.thinkwithgoogle.com/consumer-insights/consumer-trends/mobile-site-speed-importance/",
    "SENA — Servicio Nacional de Aprendizaje. (2023). Programa de formación: Análisis y Desarrollo de Software (ADSO) — Componente de infraestructura. SENA. https://www.sena.edu.co",
]

for ref in references:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(ref)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("2C2C2C")


# ════════════════════════════════════════════════════════════════
#  GUARDAR DOCUMENTO
# ════════════════════════════════════════════════════════════════

output_path = r"d:\ONEDRIVE\SENA\PROYECTOS\SWO\PLAN_DE_VALIDACION_HARDWARE_SWO.docx"
doc.save(output_path)
print(f"Documento generado exitosamente en: {output_path}")
